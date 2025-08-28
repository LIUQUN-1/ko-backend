import os
from django.http import JsonResponse
from minio import Minio, InvalidResponseError, S3Error
import hashlib
from neo4j import GraphDatabase
from bs4 import BeautifulSoup
import errno
import json
import datetime
import neo4j
import re
from random import uniform
from pdfminer.high_level import extract_text
import jieba
import jieba.analyse
from collections import Counter
import pymysql
import requests
from docx import Document
import redis
from query_neo4j.WSD import xiaoqi_instance, jiekou_3
import time
import chardet
# import pkuseg



# seg = pkuseg.pkuseg(model_name='medicine')
from requests import RequestException


class MySQLDatabase:
    def __init__(self, host, user, password, database, charset="utf8mb4"):
        """
        初始化数据库连接
        """
        self.config = {
            "host": host,
            "user": user,
            "password": password,
            "database": database,
            "charset": charset
        }
        self.connection = None

    def insert_xiaoqi_new(self, table_name, data, primary_key='xiaoqi_name'):
        try:
            existing_id = None
            # 获取主键值（支持自定义主键字段名）
            primary_key_value = data.get(primary_key)

            # 1. 检查记录是否已存在
            check_query = f"SELECT {primary_key}, xiaoqi_id FROM {table_name} WHERE {primary_key} = %s"
            with self.connection.cursor() as cursor:
                cursor.execute(check_query, (primary_key_value,))
                existing_id = cursor.fetchone()

                if existing_id:
                    print(f"主键 {primary_key_value} 已存在，返回现有ID")
                    return existing_id[1]  # 直接返回已存在记录的ID

            # 2. 执行插入操作
            columns = ", ".join(data.keys())
            placeholders = ", ".join(["%s"] * len(data))
            insert_query = f"INSERT INTO {table_name} ({columns}) VALUES ({placeholders})"

            with self.connection.cursor() as cursor:
                cursor.execute(insert_query, tuple(data.values()))
                self.connection.commit()
                return cursor.lastrowid  # 返回新插入的ID[1,4](@ref)

        except pymysql.MySQLError as e:
            print(f"插入失败：{e}")
            self.connection.rollback()

            # 3. 插入失败时再次查询已存在ID
            try:
                with self.connection.cursor() as cursor:
                    cursor.execute(check_query, (primary_key_value,))
                    existing_id = cursor.fetchone()
                    return existing_id[0] if existing_id else None
            except Exception as e:
                print(f"二次查询失败：{e}")
                return None

    def insert_data_without_primary(self, table_name, data):
        try:
            # 先检查数据是否已存在
            with self.connection.cursor() as cursor:
                # 构建WHERE条件
                conditions = " AND ".join([f"{k} = %s" for k in data.keys()])
                check_query = f"SELECT COUNT(*) FROM {table_name} WHERE {conditions}"

                cursor.execute(check_query, tuple(data.values()))
                result = cursor.fetchone()

                if result[0] > 0:
                    return False

            # 数据不存在，执行插入
            columns = ", ".join(data.keys())
            placeholders = ", ".join(["%s"] * len(data))
            insert_query = f"INSERT INTO {table_name} ({columns}) VALUES ({placeholders})"

            with self.connection.cursor() as cursor:
                cursor.execute(insert_query, tuple(data.values()))
                self.connection.commit()
                return True

        except pymysql.MySQLError as e:
            self.connection.rollback()
            return False

    def search_file(self, result, file_dict, file_dict_rev):

        # 优化数据库连接：单次连接处理所有查询
        try:
            self.connect()
            with self.connection.cursor() as cursor:
                for entity in list(result.keys()):  # 遍历副本防止迭代修改
                    file_ids = result[entity][1]
                    # 创建新列表存储有效file_id
                    valid_file_ids = []
                    for file_id in file_ids:
                        # 使用参数化查询防止SQL注入 [[5]]
                        query = "SELECT id, path FROM file WHERE id = %s"
                        cursor.execute(query, (file_id,))
                        row = cursor.fetchone()  # 假设id唯一，使用fetchone()

                        if not row:  # 未查询到结果时跳过
                            continue

                        current_path = row[1]
                        # 检查路径是否以HTTP开头 [[7]]
                        if current_path.lower().startswith(('http://', 'https://')):
                            continue  # 跳过HTTP路径

                        # 保留有效file_id和路径
                        valid_file_ids.append(file_id)
                        file_dict[int(file_id)] = current_path  # 更新路径
                        file_dict_rev[current_path] = int(file_id)
                    # 更新result中的有效file_ids列表
                    if valid_file_ids:
                        result[entity][1] = valid_file_ids
        except Exception as e:
            print(f"Error occurred: {e}")
        finally:
            self.connection.close()

        return file_dict
    def get_dir_private_list(self, entity_id, userid):
        """
        查询dir_entity，如果没结果则查询dir_to_entity -> directory，并写入新内容后返回third列表
        """

        with self.connection.cursor() as cursor:
            # 先查 dir_entity
            sql1 = """
                SELECT dir_private 
                FROM dir_entity 
                WHERE entity_id = %s AND userid = %s
            """
            cursor.execute(sql1, (entity_id, userid))
            result = cursor.fetchall()
            if result:
                return [row[0] for row in result]

            # 查不到 -> 查 dir_to_entity 获取 second
            sql2 = "SELECT second FROM dir_to_entity WHERE entity_id = %s"
            cursor.execute(sql2, (entity_id,))
            second_result = cursor.fetchone()
            if not second_result:
                return []  # 如果连 second 都查不到，直接返回空列表

            second = second_result[0]

            # 查 directory 表获取 third
            sql3 = "SELECT third FROM directory WHERE second = %s"
            cursor.execute(sql3, (second,))
            third_results = cursor.fetchall()
            third_list = [row[0] for row in third_results]

            # 插入到 dir_entity 表
            insert_sql = """
                INSERT INTO dir_entity (entity_id, dir_private, dir_sys, userid) 
                VALUES (%s, %s, %s, %s)
            """
            for third in third_list:
                cursor.execute(insert_sql, (entity_id, third, third, userid))

            self.connection.commit()

            return third_list

    def query_dir_by_name_id(self, name_id):
        """
        根据目录名称和ID查询目录
        参数:
            name_id: 目录名称和ID，格式为"名称:ID"
        返回:
            目录信息，格式为{'id': ID, 'name': 名称}，如果查询失败则返回None
        """
        try:
            with self.connection.cursor() as cursor:
                # 执行查询
                sql = "SELECT second  FROM dir_to_entity WHERE entity_id = %s"
                cursor.execute(sql, (name_id,))
                for second in cursor.fetchall():
                    return second[0]
        except pymysql.MySQLError as e:
            print(f"查询失败：{e}")
            return ""

    def insert_dir_toentity(self, classification_data, entity_id):
        try:
            if classification_data.get('code') != 200 or 'data' not in classification_data:
                print("分类数据格式不正确或API返回错误")
                return None
            data = classification_data['data']
            label_2 = data.get('label_2', '')
            # 2. 为每个文件插入记录
            file_data = {
                'second': label_2,
                'entity_id': entity_id,
            }
            self.insert_data_without_primary('dir_to_entity', file_data)
        except Exception as e:
            print(f"插入分类结果失败: {e}")
            if self.connection:
                self.connection.rollback()
            return None
    def get_directory(self, label_1, label_2):
        try:
            # 检查目录是否已存在
            query = "SELECT id FROM directory WHERE second = %s AND third = %s"
            with self.connection.cursor() as cursor:
                cursor.execute(query, (label_1, label_2))
                result = cursor.fetchone()

                if result:
                    return result[0]  # 返回已存在的目录ID

        except pymysql.MySQLError as e:
            print(f"目录插入或查询失败: {e}")
            self.connection.rollback()
            return None

    def upload_direct(self, xiaoqi, second_classify):
        # 验证输入数据有效性
        if not second_classify.get('data', {}).get('entity_path'):
            print(f"Invalid input data: {second_classify}")
            return

        original_path = second_classify['data']['entity_path']
        try:
            # 路径转换逻辑
            parts = original_path.strip('\\').split('\\')
            mapping = {'心理学': '心理'}  # 可扩展的映射规则
            processed_parts = ['KO目录'] + [mapping.get(p, p) for p in parts]
            new_entity_path = ['->'.join(processed_parts)]

            # 数据库操作
            with self.connection.cursor() as cursor:
                # 检查目录状态（注意表名保持一致性）
                query_sql = "SELECT directory FROM xiaoqi_new WHERE xiaoqi_name = %s"
                cursor.execute(query_sql, (xiaoqi,))
                result = cursor.fetchone()

                if result is None or result[0] is None:
                    # 使用JSON序列化替代字符串强转（更规范的存储方式）
                    directory_json = json.dumps(new_entity_path, ensure_ascii=False)
                    update_sql = "UPDATE xiaoqi_new SET directory = %s WHERE xiaoqi_name = %s"

                    try:
                        cursor.execute(update_sql, (directory_json, xiaoqi))
                        self.connection.commit()
                    except Exception as e:
                        self.connection.rollback()
                        print(f"❌ 更新失败: {str(e)}")
                else:
                    print(f"⚠️ 跳过更新: {xiaoqi} 已存在目录或不存在该实体")
        except Exception as e:
            print(f"❌ 处理异常: {str(e)}")
            # 可添加更详细的错误处理逻辑
    def insert_classification_result(self, classification_data, file_dict_rev, entity_id, userid):
        """
        插入分类结果到数据库
        参数:
        - classification_data: API返回的分类结果数据
        格式如: {'code': 200, 'data': {'files': {'file1.html': '类别1'}, 'label_1': '人物', 'label_2': '教育人物'}}
        """
        try:
            if classification_data.get('code') != 200 or 'data' not in classification_data:
                print("分类数据格式不正确或API返回错误")
                return None

            data = classification_data['data']
            files = data.get('files', {})
            label_1 = data.get('label_1', '')
            label_2 = data.get('label_2', '')

            results = []

            # 2. 为每个文件插入记录
            for filename, file_category in files.items():
                directory_id = self.get_directory(label_2, file_category)
                file_data = {
                    'id': directory_id,
                    'fileid': file_dict_rev[f"bb/{filename}"],
                }
                self.insert_data_without_primary('dir_to_file', file_data)
                self.get_dir_private_list(entity_id, userid)
                pri_dir_id = self.get_new_directory(entity_id, userid, file_category)
                file_data = {
                    'dir_id': pri_dir_id,
                    'file_id': file_dict_rev[f"bb/{filename}"],
                }
                self.insert_data_without_primary("dir_file", file_data)
            return results
        except Exception as e:
            print(f"插入分类结果失败: {e}")
            if self.connection:
                self.connection.rollback()
            return None

    def get_new_directory(self, entity_id, userid, label_2):
        """
        根据entity_id和userid查询dir_private字段，并以列表形式返回
        """
        with self.connection.cursor() as cursor:
            sql = """
                  SELECT id 
                  FROM dir_entity 
                  WHERE entity_id = %s AND userid = %s AND dir_private = %s
              """
            cursor.execute(sql, (entity_id, userid, label_2))
            result = cursor.fetchall()
            return result[0]
    def connect(self):
        """
        建立数据库连接
        """
        try:
            self.connection = pymysql.connect(**self.config)
            print("数据库连接成功！")
        except pymysql.MySQLError as e:
            print(f"数据库连接失败：{e}")
            raise
    def query_entities_by_name(self, name):

        entity_list = {}
        entity_dict = {}
        try:
            with self.connection.cursor() as cursor:
                # 执行模糊查询
                sql = """SELECT xiaoqi_name, key_words,xiaoqi_id 
                        FROM xiaoqi_new 
                        WHERE xiaoqi_name LIKE %s"""
                cursor.execute(sql, (f"{name}_",))


                # 处理查询结果
                for xiaoqi_name, keywords, xiaoqi_id in cursor.fetchall():
                    keywords = keywords.replace("'", '"').replace("‘", '"').replace("’", '"')
                    keyword_list = json.loads(keywords)

                    entity_list[xiaoqi_name] = keyword_list
                    entity_dict[xiaoqi_name] = xiaoqi_id
        except pymysql.MySQLError as e:
            print(f"查询失败：{e}")
            return {}

        return entity_list, entity_dict
    def insert_data(self, table_name, data):
        try:
            # 先检查主键是否存在
            primary_key = list(data.keys())[0]  # 假设主键在第一个位置
            primary_key_value = data[primary_key]

            # 生成检查主键是否存在的 SQL 查询
            check_query = f"SELECT COUNT(*) FROM {table_name} WHERE {primary_key} = %s"
            with self.connection.cursor() as cursor:
                cursor.execute(check_query, (primary_key_value,))
                result = cursor.fetchone()

                if result[0] > 0:
                    print(f"主键 {primary_key_value} 已存在，跳过插入操作。")
                    return  # 主键已存在，跳过插入操作

            # 生成插入 SQL 语句
            columns = ", ".join(data.keys())
            placeholders = ", ".join(["%s"] * len(data))
            insert_query = f"INSERT INTO {table_name} ({columns}) VALUES ({placeholders})"

            # 执行插入操作
            with self.connection.cursor() as cursor:
                cursor.execute(insert_query, tuple(data.values()))
                self.connection.commit()
                print("数据插入成功！")
        except pymysql.MySQLError as e:
            print(f"插入数据失败：{e}")
            self.connection.rollback()  # 回滚事务

    def insert_relation(self, table_name, data):
        try:
            # 先检查主键是否存在
            primary_key = list(data.keys())[0]  # 假设主键在第一个位置
            last_key = list(data.keys())[-1]
            primary_key_value = data[primary_key]
            last_key_value = data[last_key]

            # 生成检查主键是否存在的 SQL 查询
            check_query = f"SELECT COUNT(*) FROM {table_name} WHERE {primary_key} = %s AND {last_key} = %s "
            with self.connection.cursor() as cursor:
                cursor.execute(check_query, (primary_key_value, last_key_value, ))
                result = cursor.fetchone()

                if result[0] > 0:
                    print(f"主键 {primary_key_value} 已存在，跳过插入操作。")
                    return  # 主键已存在，跳过插入操作

            # 生成插入 SQL 语句
            columns = ", ".join(data.keys())
            placeholders = ", ".join(["%s"] * len(data))
            insert_query = f"INSERT INTO {table_name} ({columns}) VALUES ({placeholders})"

            # 执行插入操作
            with self.connection.cursor() as cursor:
                cursor.execute(insert_query, tuple(data.values()))
                self.connection.commit()
                print("数据插入成功！")
        except pymysql.MySQLError as e:
            print(f"插入数据失败：{e}")
            self.connection.rollback()  # 回滚事务

    def close(self):
        """
        关闭数据库连接
        """
        if self.connection:
            self.connection.close()
            print("数据库连接已关闭！")


def Dir_html_word(html_file_path,flag):
    try:
        # 打开并读取本地HTML文件
        html_content = None
        with open(html_file_path, 'rb') as f:
            raw_data = f.read()

            detected = chardet.detect(raw_data)
            encoding = detected['encoding'] if detected['confidence'] > 0.7 else 'utf-8'

            try:
                html_content = raw_data.decode(encoding, errors='replace')
            except (LookupError, UnicodeDecodeError):
                # 如果检测的编码无效，尝试常见中文编码
                html_content = raw_data.decode('gbk', errors='replace')

        # 使用BeautifulSoup解析HTML内容
        soup = BeautifulSoup(html_content, 'html.parser')

        main_content = []
        image_content = []


        if flag == 1:#维基百科
            # 假设我们想要提取所有的段落内容，即<p>标签
            paragraphs = soup.find_all('p')

            # 遍历所有的段落，提取文本
            for p in paragraphs:
                # .strip()用于去除字符串首尾的空白字符和\n
                main_content.append(p.get_text().strip())

            meta_img = soup.find_all('meta')
            for tag in meta_img:
                if 'property' in tag.attrs and tag['property'].lower() == 'og:image':
                    content = tag.get('content')
                    image_content.append(content)


        elif flag == 2:#百度百科
            # 查找所有的meta标签
            meta_tags = soup.find_all('meta')
            for tag in meta_tags:
                if 'name' in tag.attrs and tag['name'].lower() == 'description':
                    content = tag.get('content')
                    main_content.append(content)
                elif 'property' in tag.attrs and tag['property'].lower() == 'og:description':
                    content = tag.get('content')
                    main_content.append(content)
                if 'name' in tag.attrs and tag['name'].lower() == 'image':
                    content = tag.get('content')
                    image_content.append(content)
            #查找所有class为text_bypwF的<span>标签
            span_tags = soup.find_all('span', class_='text_bypwF')
            str_innner = ""
            # 遍历这些标签并打印它们的内容
            for tag in span_tags:
                # 如果<span>标签内部有<a>标签，需要进一步提取<a>标签的内容
                if tag.find('a'):
                    # 提取<a>标签的文本内容
                    link_text = tag.find('a').get_text()
                    str_innner += link_text
                else:
                    # 直接提取<span>标签的文本内容
                    str_innner += tag.get_text()
            main_content.append(str_innner)

        # 返回提取的主要内容
        return main_content

    except FileNotFoundError:
        print("文件未找到，请检查文件路径是否正确。")
    except IOError as e:
        print(f"文件读取错误: {e}")
    except Exception as e:
        print(f"其他错误: {e}")

class Bucket:

    def __init__(self, minio_address, minio_admin, minio_password):
        # 通过ip 账号 密码 连接minio server
        # Http连接 将secure设置为False
        self.minioClient = Minio(endpoint=minio_address,
                                 access_key=minio_admin,
                                 secret_key=minio_password,
                                 secure=False)

    def create_one_bucket(self, bucket_name):
        # 创建桶(调用make_bucket api来创建一个桶)
        """
        桶命名规则：小写字母，句点，连字符和数字 允许使用 长度至少3个字符
        使用大写字母、下划线等会报错
        """
        try:
            # bucket_exists：检查桶是否存在
            if self.minioClient.bucket_exists(bucket_name=bucket_name):
                print("该存储桶已经存在")
            else:
                self.minioClient.make_bucket(bucket_name=bucket_name)
                print(f"{bucket_name}桶创建成功")
        except InvalidResponseError as err:
            print(err)

    def remove_one_bucket(self, bucket_name):
        # 删除桶(调用remove_bucket api来创建一个存储桶)
        try:
            if self.minioClient.bucket_exists(bucket_name=bucket_name):
                self.minioClient.remove_bucket(bucket_name)
                print("删除存储桶成功")
            else:
                print("该存储桶不存在")
        except InvalidResponseError as err:
            print(err)


    def upload_file_to_bucket(self, bucket_name, file_name, file_path):
        """
        将文件上传到bucket
        :param bucket_name: minio桶名称
        :param file_name: 存放到minio桶中的文件名字(相当于对文件进行了重命名，可以与原文件名不同)
                            file_name处可以创建新的目录(文件夹) 例如 /example/file_name
                            相当于在该桶中新建了一个example文件夹 并把文件放在其中
        :param file_path: 本地文件的路径
        """
        # 桶是否存在 不存在则新建
        check_bucket = self.minioClient.bucket_exists(bucket_name)
        if not check_bucket:
            self.minioClient.make_bucket(bucket_name)

        try:
            self.minioClient.fput_object(bucket_name=bucket_name,
                                         object_name=file_name,
                                         file_path=file_path)
        except FileNotFoundError as err:
            print('upload_failed: ' + str(err))
        except S3Error as err:
            print("upload_failed:", err)

    def download_file_from_bucket(self, bucket_name, minio_file_path, download_file_path):
        """
        从bucket下载文件
        :param bucket_name: minio桶名称
        :param minio_file_path: 存放在minio桶中文件名字
                            file_name处可以包含目录(文件夹) 例如 /example/file_name
        :param download_file_path: 文件获取后存放的路径
        """
        # 桶是否存在
        check_bucket = self.minioClient.bucket_exists(bucket_name)
        if check_bucket:
            try:
                self.minioClient.fget_object(bucket_name=bucket_name,
                                             object_name=minio_file_path,
                                             file_path=download_file_path)
                return 1
            except FileNotFoundError as err:
                print('download_failed: ' + str(err))
                return 0
            except S3Error as err:
                print("download_failed:", err)

    def remove_object(self, bucket_name, object_name):
        """
        从bucket删除文件
        :param bucket_name: minio桶名称
        :param object_name: 存放在minio桶中的文件名字
                            object_name处可以包含目录(文件夹) 例如 /example/file_name
        """
        # 桶是否存在
        check_bucket = self.minioClient.bucket_exists(bucket_name)
        if check_bucket:
            try:
                self.minioClient.remove_object(bucket_name=bucket_name,
                                               object_name=object_name)
            except FileNotFoundError as err:
                print('upload_failed: ' + str(err))
            except S3Error as err:
                print("upload_failed:", err)

    # 获取所有的桶
    def get_all_bucket(self):
        buckets = self.minioClient.list_buckets()
        ret = []
        for _ in buckets:
            ret.append(_.name)
        return ret

    # 获取一个桶中的所有一级目录和文件
    def get_list_objects_from_bucket(self, bucket_name):
        # 桶是否存在
        check_bucket = self.minioClient.bucket_exists(bucket_name)
        if check_bucket:
            # 获取到该桶中的所有目录和文件
            objects = self.minioClient.list_objects(bucket_name=bucket_name)
            ret = []
            for _ in objects:
                ret.append(_.object_name)
            return ret

    # 获取桶里某个目录下的所有目录和文件
    def get_list_objects_from_bucket_dir(self, bucket_name, dir_name):
        # 桶是否存在
        check_bucket = self.minioClient.bucket_exists(bucket_name)
        if check_bucket:
            # 获取到bucket_所name桶中的dir_name下的有目录和文件
            # prefix 获取的文件路径需包含该前缀
            objects = self.minioClient.list_objects(bucket_name=bucket_name,
                                                    prefix=dir_name,
                                                    recursive=True)
            ret = []
            for obj in objects:
                object_name = obj.object_name
                # 获取对象的内容
                content = self.minioClient.get_object(bucket_name=bucket_name,
                                                      object_name=object_name)
                ret.append(content.data.decode())
            return ret

def add_to_mysql(result, entity, sim, id):
    db = MySQLDatabase(
        host="114.213.234.179",
        user="koroot",  # 替换为您的用户名
        password="DMiC-4092",  # 替换为您的密码
        database="db_hp"  # 替换为您的数据库名
    )
    try:
        db.connect()
        data_to_insert = {
            "id": str(id),  # 假设 id 是自增字段，可以设置为 None
            "name": result[0]['h']["name"],  # 替换为实际数据
            "path": result[0]['h']["path"],  # 替换为实际时间戳，格式：YYYY-MM-DD HH:MM:SS
            "timestamp": result[0]['h']["timestamp"],  # 替换为实际 URL
            "private": 1,
            "userid": result[0]['h']["user_id"]
        }
        relation_to_insert = {
            "entity": entity,
            "sim": sim,
            "file_id": str(id)
        }
        db.insert_data("File", data_to_insert)
        db.insert_relation("entity_to_file",relation_to_insert)
    finally:
        # 关闭数据库连接
        db.close()

def read_txt_file(file_path):
    with open(file_path, 'r', encoding="utf-8") as file:
        file_contents = file.read()
    return file_contents
def get_sha1_hash(file_name):
    shal_hash = hashlib.sha1(file_name.encode()).hexdigest()
    return shal_hash


def create_entity_and_link(tx, name, path, userid):
    import datetime
    current_date = datetime.datetime.now()
    year = current_date.year
    month = current_date.month
    day = current_date.day
    minute = current_date.minute
    formatted_date = f"{year}-{month}-{day}"
    formatted_min = f"{year}-{month}-{day}_{minute}"

    query = """
            MERGE (f:Strict {name: $filename, path: $path, private: $private, user_id: $user_id, timestamp: $time})
            WITH f
            RETURN id(f) as id
            """
    # RETURN id(f) AS file_id, id(h) AS hypernode_id, id(k) AS category_id
    result = tx.run(
        query,
        filename=name,
        path=path,
        private=1,
        user_id=userid,
        time=formatted_date,
    )
    record = result.single()
    return record["id"]
# def create_entity_and_link(tx, hypename, filename, path, catename,userID):
#     import datetime
#     current_date = datetime.datetime.now()
#     year = current_date.year
#     month = current_date.month
#     day = current_date.day
#     minute = current_date.minute
#     url = "ko.zhonghuapu.com/hypernode/" + hypename
#     formatted_date = f"{year}-{month}-{day}"
#     formatted_min = f"{year}-{month}-{day}_{minute}"
#     name = filename.split('.')[0] + '_' + formatted_min + '.' + filename.split('.')[-1]
#
#     query = """
#             MERGE (f:Strict {name: $filename, path: $path, private: $private, user_id: $user_id, timestamp: $time})
#             WITH f
#             MERGE (h:test {url: $url, name: $hypename})
#             SET h.timestamp = $time, h.type = 'hypernode'
#             WITH f, h
#             MERGE (h)-[:hyperedge]->(f)
#             WITH f, h
#             MATCH (cat:KOCategory {name: $cate_name})
#             MERGE (cat)-[c:edge]->(h)
#             RETURN h, c, id(f) as id
#             """
#     # RETURN id(f) AS file_id, id(h) AS hypernode_id, id(k) AS category_id
#     result = tx.run(
#         query,
#         filename=name,
#         hypename=hypename,
#         path=path,
#         private=1,
#         user_id=int(userID),
#         time=formatted_date,
#         url=url,
#         cate_name=catename
#     )
#     record = result.single()
#     return record["id"], name
def redis_filter(list):
    r = redis.Redis(
        host='114.213.232.140',
        port=16379,
        decode_responses=True
    )
    start = time.time()
    # 测试连接

    def check_values_exist(redis_client, key, values):
        pipe = redis_client.pipeline()
        for value in values:
            pipe.sismember(key, value)  # 将多个 SISMEMBER 命令加入管道 <button class="citation-flag" data-index="10">
        results = pipe.execute()  # 批量执行命令并获取结果
        return dict(zip(values, results))  # 将结果映射为字典

    # 执行检查
    existence_dict = check_values_exist(r, 'kocategory_names', list)
    true_keys = [k for k in list if existence_dict[k]]
    end = time.time()
    print(f"共耗时{end-start}秒")
    return true_keys

def create_entity_node(tx, hypename ,filename, path, userID):
    current_date = datetime.datetime.now()
    year = current_date.year
    month = current_date.month
    day = current_date.day
    minute = current_date.minute
    url = "ko.zhonghuapu.com/hypernode/"+hypename
    formatted_date = f"{year}-{month}-{day}"
    formatted_min = f"{year}-{month}-{day}_{minute}"
    name = filename.split('.')[0]+'_'+formatted_min+'.'+filename.split('.')[-1]
    # query = f"MERGE (n:File {{name: $name}}) \
    # ON CREATE SET n.path = $path, n.private = $private, n.user_id = $id, n.timestamp=$time RETURN n \
    query = f"MERGE (f:Strict {{name: $filename, path: $path, private: $private, user_id: $id, timestamp: $time}})\
    WITH f\
    MERGE (h:hypernode {{url: $url}})  \
    SET h.name = $hypename, h.timestamp = $time, h.type = 'hypernode'\
    WITH f, h\
    MERGE (h)-[:hyperedge]->(f)\
    RETURN id(f) as id"
    result = tx.run(query, filename=name, hypename=hypename, path=path, private=1, id=int(userID), time=formatted_date, url=url)
    return [record['id'] for record in result][0], name

def create_KoCate_link(tx, name,cate_name):
    query = f"MATCH (f:KOCategory {{name: $cate_name}}) \
    MATCH (h:hypernode {{name: $name}}) \
    MERGE  (f)-[c:hyperedge]->(h) \
    RETURN h,c, f"
    result = tx.run(query, cate_name=cate_name, name=name)

    return result

def create_link(tx, name):
    query = f"MATCH (f:KOcategory {{name: 'KO目录'}}) \
    MATCH (h:hypernode {{name: $name}}) \
    MERGE  (f)-[:hyperedge]->(h) \
    RETURN h, f"
    tx.run(query, name = name)
def find_common_characters(str1, str2):
    set1 = set(jieba.lcut(str1))
    set2 = set(jieba.lcut(str2))
    common_chars = set1.intersection(set2)
    return list(common_chars)
def get_entity_node(tx, name):
    query = (f"MATCH (h:hypernode)<-[r:Superedge]-(t) WHERE (h.name=$name and t:wikibaike) or (h.name=$name and t:baidu_directory)\
             RETURN t")
    result = tx.run(query, name= name)
    return result.data()

def create_strict_and_coaser(tx, id, direct):

    query = """
        MATCH (a:Strict) WHERE id(a) = $id    
        MATCH (b:coarse {name: $direct})       
        MERGE (a)-[r:edge]->(b)
         RETURN r IS NOT NULL AS exists   
        """
    # RETURN id(f) AS file_id, id(h) AS hypernode_id, id(k) AS category_id
    result = tx.run(
        query,
        id = id,
        direct = direct
    )
    return result.single()["exists"]
def search_name_contain(tx, name):
    query_word = ("MATCH (h:hypernode) WHERE h.name CONTAINS $name RETURN h.name as name")
    result = tx.run(query_word, name=name)
    return result.data()
def search_name_is(tx, name):
    query_word = ("MATCH (h:hypernode) WHERE h.name = $name RETURN h.name as name")
    result = tx.run(query_word, name=name)
    return result.data()
def search_File_name(tx, name):
    query_word = ("MATCH (h:Strict) WHERE h.name = $name RETURN h")
    result = tx.run(query_word, name=name)
    return result.data()
def search_file(tx, name):
    query_word = ("MATCH (h:hypernode)-[r]-(t) WHERE h.name = $name AND (t:wikipage OR t:baidupage OR t:File) RETURN COALESCE(t.file_path, t.path) as file_path, t.name as name, labels(t) as label, h.name as hname")
    result = tx.run(query_word, name=name)
    return result.data()
def find_common_characters(str1, str2):
    set1 = set(jieba.lcut(str1))
    set2 = set(jieba.lcut(str2))
    common_chars = set1.intersection(set2)
    return list(common_chars)
def extract_text_from_docx(docx_path):
    try:
        # 打开docx文件
        doc = Document(docx_path)

        # 提取所有段落文本
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)

        # 提取表格中的文本
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)

        # 合并所有文本，用换行符分隔
        return '\n'.join(full_text)

    except Exception as e:
        print(f"Error processing {docx_path}: {str(e)}")
        return None
def push_file_info_mysql(session,words,stop_words,file_id,file_name ,file_dict, file_dict_rev,userID):
    filtered_words = [word for word in words if word not in stop_words and word.strip()]
    filtered_words = redis_filter(filtered_words)
    word_counts = Counter(filtered_words)
    sorted_word_counts = word_counts.most_common()
    result = session.write_transaction(search_File_name, file_name)
    # add_to_mysql(result, file_obj.name.split('.')[0], 1, file_id)
    # add_to_xiaoqi(doc, file_id, file_obj.name.split('.')[0], file_dict, file_dict_rev, userID)
    for word, weight in sorted_word_counts:
        add_to_mysql(result, word, weight / len(filtered_words), file_id)
        # add_to_xiaoqi(doc, file_id, word, file_dict, file_dict_rev, userID)

def redis_filter(list):
    r = redis.Redis(
        host='114.213.232.140',
        port=16379,
        decode_responses=True
    )
    start = time.time()
    # 测试连接

    def check_values_exist(redis_client, key, values):
        pipe = redis_client.pipeline()
        for value in values:
            pipe.sismember(key, value)  # 将多个 SISMEMBER 命令加入管道 <button class="citation-flag" data-index="10">
        results = pipe.execute()  # 批量执行命令并获取结果
        return dict(zip(values, results))  # 将结果映射为字典

    # 执行检查
    existence_dict = check_values_exist(r, 'kocategory_names', list)
    true_keys = [k for k, v in existence_dict.items() if v ]
    end = time.time()
    return true_keys

def generate_entity_json(result, file_dict, info, customize_conte):
    output_list = []

    for entity in result:
        # 获取文件ID列表
        file_ids = result[entity][1]

        # 生成minio文件路径列表（网页1、网页2的类型转换思路）
        minio_paths = [file_dict[int(file_id)] for file_id in file_ids]

        # 提取关键词和关联实体（网页3的字典转换思路）
        keywords = result[entity][0]
        related_files = [file_dict[int(file_id)].split('_')[0] for file_id in file_ids]  # 网页8的字符串分割技巧
        merged_keywords = list({k: None for k in keywords + related_files}.keys())  # 网页7的去重方法

        # 构建最终结构（网页6的JSON转换技巧）
        output = {
            "minio_file_path_list": minio_paths,
            "entity": entity,
            "entity_with_keyword": f"{entity}: {', '.join(merged_keywords)}",
            "info": info if info is not None else "",
            "customize_conte": customize_conte
        }
        output_list.append(output)

    return output_list

def entity_classification(payload):
    # 接口地址（根据实际环境替换）
    url = "http://114.213.232.140:8000/api/classify/entity/"

    # 请求头配置（网页3、网页6的header设置方法）
    headers = {
        "Content-Type": "application/json",
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) PythonClient/1.0"  # 网页8的请求头伪装技巧
    }

    try:
        # 发送POST请求（网页1、网页7的核心方法）
        response = requests.post(
            url,
            headers=headers,
            json=payload,  # 网页6推荐的json参数自动序列化
            timeout=10  # 网页9的超时设置
        )

        # 状态码检查（网页4的错误处理机制）
        if response.status_code == 200:
            # 解析JSON响应（网页3、网页7的响应处理方法）
            result = response.json()
            print(json.dumps(result, indent=2, ensure_ascii=False))

            # 此处可添加业务逻辑处理
            # 例如：result.get("classification_result")

            return result
        else:
            print(f"接口调用失败，状态码：{response.status_code}")
            print(f"错误详情：{response.text}")
            return False

    except RequestException as e:  # 网页4的异常捕获机制
        print(f"请求异常：{str(e)}")
        return False
    except json.JSONDecodeError:  # 网页7的JSON解析异常处理
        print("响应数据解析失败")
        return False


def add_to_xiaoqi(file_text, file_id, entity, file_dict, file_dict_rev, userID):
    db = MySQLDatabase(
        host="114.213.234.179",
        user="koroot",
        password="DMiC-4092",
        database="db_hp"
    )
    result = {}
    customize_content = []
    driver = GraphDatabase.driver("bolt://114.213.232.140:37687", auth=("neo4j", "123456"))
    db.connect()
    entity_list, entity_dict = db.query_entities_by_name(entity)
    db.close()
    info = None
    if len(entity_list) != 0:
        dictRes = xiaoqi_instance(file_text, entity_list)
        db.connect()
        info = db.query_dir_by_name_id(entity_dict[dictRes["entity"]])
        customize_content = db.get_dir_private_list(dictRes["entity"], userID)
        result[dictRes["entity"]] = [dictRes["entity_with_keyword"].split(','), [file_id]]
        entity_dict = updata_to_mysql_new(result, True)
    else:
        file_data = {
            'entity': entity,
            'sim': 1,
            'file_id': file_id,
        }
        db.connect()
        db.insert_data_without_primary("entity_to_file", file_data)
        result, _, _ = jiekou_3(entity, userID)
        entity_dict = updata_to_mysql_new(result, True)
        file_dict = db.search_file(result, file_dict, file_dict_rev)
    final_output = generate_entity_json(result, file_dict, info, customize_content)
    for payload in final_output:
        first_classify = first_classification(payload)
        second_classify = second_classification(payload)
        try:
            db.connect()
            db.insert_dir_toentity(first_classify, entity_dict[payload["entity"]])
            db.insert_classification_result(first_classify, file_dict_rev, entity_dict[payload["entity"]], userID)
            db.upload_direct(payload["entity"], second_classify)
            with driver.session() as session:
                direct = second_classify['data']['file_path'][file_dict[file_id].split('/')[-1]]
                result = session.write_transaction(
                    create_strict_and_coaser,
                    id=file_id,
                    direct = direct if direct != "Root" else "KO目录" ,
                )
                if result == False:
                    print("上传失败 error")
        finally:
            # 关闭数据库连接
            db.close()


def first_classification(payload):
    # 接口地址（根据实际环境替换）
    url = "http://114.213.232.140:8000/api/classify/entity/"

    # 请求头配置（网页3、网页6的header设置方法）
    headers = {
        "Content-Type": "application/json",
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) PythonClient/1.0"  # 网页8的请求头伪装技巧
    }

    try:
        # 发送POST请求（网页1、网页7的核心方法）
        response = requests.post(
            url,
            headers=headers,
            json=payload,  # 网页6推荐的json参数自动序列化
            # timeout=10  # 网页9的超时设置
        )

        # 状态码检查（网页4的错误处理机制）
        if response.status_code == 200:
            # 解析JSON响应（网页3、网页7的响应处理方法）
            result = response.json()

            # 此处可添加业务逻辑处理
            # 例如：result.get("classification_result")

            return result
        else:
            print(f"接口调用失败，状态码：{response.status_code}")
            print(f"错误详情：{response.text}")
            return False

    except json.JSONDecodeError:  # 网页7的JSON解析异常处理
        print("响应数据解析失败")
        return False


def second_classification(payload):
    # 接口地址（根据实际环境替换）
    url = "http://114.213.232.140:8000/api/classify/big-content_last/"

    # 请求头配置（网页3、网页6的header设置方法）
    headers = {
        "Content-Type": "application/json",
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) PythonClient/1.0"  # 网页8的请求头伪装技巧
    }
    try:
        # 发送POST请求（网页1、网页7的核心方法）
        response = requests.post(
            url,
            headers=headers,
            json=payload,  # 网页6推荐的json参数自动序列化
            # timeout=30  # 网页9的超时设置
        )

        # 状态码检查（网页4的错误处理机制）
        if response.status_code == 200:
            # 解析JSON响应（网页3、网页7的响应处理方法）
            result = response.json()

            return result
        else:
            print(f"接口调用失败，状态码：{response.status_code}")
            print(f"错误详情：{response.text}")
            return False
    except json.JSONDecodeError:  # 网页7的JSON解析异常处理
        print("响应数据解析失败")
        return False

def updata_to_mysql_new(result, is_xiaoqi):
    db = MySQLDatabase(
        host="114.213.234.179",
        user="koroot",  # 替换为您的用户名
        password="DMiC-4092",  # 替换为您的密码
        database="db_hp"  # 替换为您的数据库名
    )
    entity_file = {}
    db.connect()
    try:
        for key, value in result.items():
            if is_xiaoqi:
                data_to_insert = {
                    "key_words": str(value[0]),
                    "xiaoqi_name": str(key)
                }
                # print(key)
                entity_file[key] = db.insert_xiaoqi_new("xiaoqi_new", data_to_insert)
            for i in value[1]:
                query = """
                            INSERT INTO xiaoqi_to_file (xiaoqi_id, file_id) 
                SELECT 
                    xnew.xiaoqi_id, %s 
                FROM 
                    xiaoqi_new xnew 
                WHERE 
                    xnew.xiaoqi_name = %s
                LIMIT 1; 
                            """
                try:
                    with db.connection.cursor() as cursor:
                        cursor.execute(query, (int(i), key))
                        db.connection.commit()
                        # 检查受影响的行数
                        if cursor.rowcount > 0:
                            print(f'更新成功，受影响的行数: {cursor.rowcount}')
                        else:
                            print('没有找到匹配的行，更新未执行。')
                except pymysql.MySQLError as e:
                    print(f"查询失败：{e}")
                    raise

    finally:
        db.close()
        # 关闭数据库连接
        return entity_file
def main(request):
    strict=str(request)
    strict=strict.replace("<","")
    strict = strict.replace(">", "")
    strict = strict.replace("\'", "")
    strict = strict.split("?strict=")
    userID = strict[1]
    strict = strict[1][0]
    userID = userID.split("&userID=")
    userID = userID[1]
    # strict = request.args.get('strict', default='0', type=str)
    jieba.analyse.set_stop_words("G:\KoDjango\hit_stopwords\hit_stopwords.txt")
    stop_words = set()
    with open(r"G:\KoDjango\hit_stopwords\hit_stopwords.txt", "r", encoding="utf-8") as f:
        stop_words.update(line.strip() for line in f)
    minio_address = "114.213.232.140:19000"
    minio_admin = "minioadmin"
    minio_password = "minioadmin"
    bucket = Bucket(minio_address=minio_address,
                    minio_admin=minio_admin,
                    minio_password=minio_password)
    # 创建桶测试
    bucket.create_one_bucket('kofiles')
    dict1 = 'F:\data\\tmp\\'
    doc = ""
    file_obj = request.FILES.get('file', None)
    head_path = 'D:/upload/'
    file_path = os.path.join(head_path, file_obj.name)
    with open(file_path, 'wb') as f:
        for chunk in file_obj.chunks():
            f.write(chunk)
    doc_type = file_obj.name.split('.')[-1]
    current_date = datetime.datetime.now()
    year = current_date.year
    month = current_date.month
    day = current_date.day
    minute = current_date.minute
    formatted_min = f"{year}-{month}-{day}_{minute}"
    file_name = file_obj.name.split('.')[0] + '_' + formatted_min + '.' + \
                file_obj.name.split('.')[-1]
    path = get_sha1_hash('upload')[:2] + '/' + file_name
    contents = []
    data = []
    try:
        bucket.upload_file_to_bucket('kofiles', file_name, file_path)
        with open(file_path, 'wb') as f:
            for chunk in file_obj.chunks():
                f.write(chunk)
        if doc_type == "txt":
            doc = read_txt_file(file_path)
        elif doc_type == "pdf":
            doc = extract_text(file_path)
        elif doc_type == "html":
            doc = Dir_html_word(file_path, 2)[0]
        elif doc_type == "docx":
            doc = extract_text_from_docx(file_path)
    except neo4j.exceptions.ServiceUnavailable as e:
        error_message = {
            "error": str(e),
            "message": '上传文件失败'
        }
        print(str(e))
        return json.dumps(error_message)
    # try:
    driver = GraphDatabase.driver("bolt://114.213.232.140:37687", auth=("neo4j", "123456"))
    # driver = GraphDatabase.driver("bolt://114.213.232.140:7687", auth=("neo4j", "DMiChao"))
    with driver.session() as session:
        # result = session.write_transaction(search_name_contain, file_obj.name.split('.')[0])
        words = jieba.lcut(doc)
        filtered_words = [word for word in words if word not in stop_words and word.strip()]
        filtered_words = redis_filter(filtered_words)
        word_counts = Counter(filtered_words)
        sorted_word_counts = word_counts.most_common()
        # top_20_percent = sorted_word_counts[:max(1, int(len(sorted_word_counts) * 0.2))]
        # words = jieba.lcut(doc)
        # filtered_words = [word for word in words if word not in stop_words and word.strip()]
        # word_counts = Counter(filtered_words)
        # sorted_word_counts = word_counts.most_common()
        # top_20_percent = sorted_word_counts[:max(1, int(len(sorted_word_counts) * 0.2))]
        # worf_topk = jieba.analyse.extract_tags(doc, withWeight=True)
        # top_20_percent = worf_topk[:max(1, int(len(worf_topk) * 0.2))]
        # 发送 GET 请求
        # payload = {
        #     'token': doc  # 将 doc[:3000] 作为 payload 传递
        # }
        # url = "http://114.213.232.140:18090/KoCategory/"
        # try:
        #     # 发送 GET 请求
        #     response = requests.post(url, json=payload)
        #     # 检查 HTTP 响应状态
        #     if response.status_code == 200:
        #         data = response.json()  # 获取 JSON 数据
        #         print("Response from Django API:", data)
        #     else:
        #         print("Failed to call Django API. Status code:", response.status_code)
        # except requests.exceptions.RequestException as e:
        #     print("Error occurred while calling Django API:", e)
        # file_id, file_name = session.write_transaction(create_entity_node, file_obj.name.split('.')[0],
        #                                                file_obj.name,
        #                                                file_name,
        #                                                data[-1])
        # session.write_transaction(create_KoCate_link, file_obj.name.split('.')[0], data[-1])
        if len(data) <= 0:
            data = ["KO目录"]
        file_id = session.write_transaction(
            create_entity_and_link,
            file_name,
            path,
            userID
        )
        session.close()
        # 将文件名作为实体将文件名与文件的关系上传到服务器
        result = session.write_transaction(search_File_name, file_name)
        # file_dict = {file_id: file_name}
        # file_dict_rev = {file_name: file_id}
        # add_to_xiaoqi(doc, file_id, file_obj.name.split('.')[0], file_dict, file_dict_rev, userID)
        for word, weight in sorted_word_counts:
            add_to_mysql(result, word, weight / len(filtered_words), file_id)
            # add_to_xiaoqi(doc, file_id, word, file_dict, file_dict_rev,userID)

    driver.close()
    message = {
        "message": "上传成功",
    }
    return json.dumps(message)