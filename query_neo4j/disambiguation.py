import os
import time
import chardet
from django.http import JsonResponse
from minio import Minio, InvalidResponseError, S3Error
import hashlib
from neo4j import GraphDatabase
# from bs4 import BeautifulSoup  # 已替换为html.parser
from html.parser import HTMLParser
import html
import errno
import json
import datetime
import neo4j
import re
from random import uniform
from pdfminer.high_level import extract_text
import jieba
import jieba.analyse
from collections import Counter
import pymysql
import requests
from docx import Document
import redis
from query_neo4j.WSD import xiaoqi_instance, jiekou_3
import fitz  # PyMuPDF
import threading
import logging
from markitdown import MarkItDown
from concurrent.futures import ThreadPoolExecutor, as_completed

# 获取Django配置的日志记录器
logger = logging.getLogger('query_neo4j')

# Redis分布式锁配置 - 参考upload.py的配置
REDIS_HOST = '114.213.232.140'  # Redis服务器地址
REDIS_PORT = 26379  # Redis端口
REDIS_DB = 0  # Redis数据库编号
REDIS_PASSWORD = None  # Redis密码


class RedisDistributedLock:
    """Redis分布式锁实现"""

    def __init__(self, redis_client, key, timeout=30, retry_times=50, retry_delay=0.1):
        """
        初始化分布式锁
        :param redis_client: Redis客户端
        :param key: 锁的键名
        :param timeout: 锁的超时时间（秒）
        :param retry_times: 重试次数
        :param retry_delay: 重试间隔（秒）
        """
        self.redis_client = redis_client
        self.key = f"distributed_lock:{key}"
        self.timeout = timeout
        self.retry_times = retry_times
        self.retry_delay = retry_delay
        self.identifier = None

    def acquire(self, blocking=True):
        """获取锁
        :param blocking: 是否阻塞等待，True为阻塞直到获取到锁，False为非阻塞模式
        """
        import uuid
        import time

        identifier = str(uuid.uuid4())

        if blocking:
            # 阻塞模式：使用Redis原生的阻塞机制
            while True:
                # 尝试获取锁
                if self.redis_client.set(self.key, identifier, nx=True, ex=self.timeout):
                    self.identifier = identifier
                    return True

                # 使用Redis的BLPOP实现阻塞等待
                # 创建一个等待队列，当锁释放时会有通知
                wait_key = f"{self.key}:wait"
                try:
                    # BLPOP会阻塞等待，直到有元素或超时
                    # 设置较短的超时(1秒)以便定期检查锁状态
                    result = self.redis_client.blpop(wait_key, timeout=2)
                    # 无论是否有通知，都再次尝试获取锁
                    continue
                except Exception:
                    # 如果BLPOP出现异常，降级为短暂sleep
                    time.sleep(0.1)
                    continue
        else:
            # 非阻塞模式：原有逻辑
            end_time = time.time() + self.timeout

            for _ in range(self.retry_times):
                if time.time() > end_time:
                    return False

                # 尝试获取锁
                if self.redis_client.set(self.key, identifier, nx=True, ex=self.timeout):
                    self.identifier = identifier
                    return True

                time.sleep(self.retry_delay)

            return False

    def release(self):
        """释放锁"""
        if not self.identifier:
            return False

        # 使用Lua脚本确保原子性，并在释放锁后通知等待的线程
        lua_script = """
        if redis.call('GET', KEYS[1]) == ARGV[1] then
            local result = redis.call('DEL', KEYS[1])
            if result == 1 then
                -- 锁释放成功，通知等待的线程
                -- 向等待队列推送一个通知信号
                redis.call('LPUSH', KEYS[2], '1')
                -- 设置通知的过期时间，避免累积过多通知
                redis.call('EXPIRE', KEYS[2], 5)
            end
            return result
        else
            return 0
        end
        """

        try:
            wait_key = f"{self.key}:wait"
            result = self.redis_client.eval(lua_script, 2, self.key, wait_key, self.identifier)
            return result == 1
        except Exception as e:
            logger.error(f"释放锁失败: {e}")
            return False

    def __enter__(self):
        """支持with语句，默认使用阻塞模式"""
        self.acquire(blocking=True)  # 阻塞等待获取锁
        return self

    def __exit__(self, exc_type, exc_val, exc_tb):
        """支持with语句"""
        self.release()


def get_redis_client():
    """获取Redis客户端"""
    try:
        client = redis.Redis(
            host=REDIS_HOST,
            port=REDIS_PORT,
            db=REDIS_DB,
            password=REDIS_PASSWORD,
            decode_responses=True
        )
        # 测试连接
        client.ping()
        return client
    except Exception as e:
        logger.error(f"Redis连接失败: {e}")
        return None


def get_distributed_lock(entity, timeout=30):
    """获取分布式锁"""
    redis_client = get_redis_client()
    if not redis_client:
        raise Exception("Redis连接失败，无法创建分布式锁")

    return RedisDistributedLock(redis_client, f"xiaoqi:{entity}", timeout)


class MySQLDatabase:
    def __init__(self, host, user, password, database, charset="utf8mb4"):
        """
        初始化数据库连接
        """
        self.config = {
            "host": host,
            "user": user,
            "password": password,
            "database": database,
            "charset": charset
        }
        self.connection = None

    def insert_data_without_primary(self, table_name, data):
        """不包含主键的数据插入（无事务控制版本），依赖外层事务管理"""
        try:
            with self.connection.cursor() as cursor:
                # 构建 WHERE 条件
                conditions = " AND ".join([f"{k} = %s" for k in data.keys()])
                check_query = f"SELECT COUNT(*) FROM {table_name} WHERE {conditions} FOR UPDATE"

                # 执行检查并加锁
                cursor.execute(check_query, tuple(data.values()))
                result = cursor.fetchone()

                if result[0] > 0:
                    logger.info("数据已存在，跳过插入")
                    return False

                # 执行插入操作
                columns = ", ".join(data.keys())
                placeholders = ", ".join(["%s"] * len(data))
                insert_query = f"INSERT INTO {table_name} ({columns}) VALUES ({placeholders})"
                cursor.execute(insert_query, tuple(data.values()))
                self.connection.commit()
                return True

        except pymysql.MySQLError as e:
            logger.error(f"插入操作失败：{e}")

    def search_file(self, result, file_dict, file_dict_rev):

        # 优化数据库连接：单次连接处理所有查询
        try:
            self.connect()
            with self.connection.cursor() as cursor:
                for entity in list(result.keys()):  # 遍历副本防止迭代修改
                    file_ids = result[entity][1]
                    # 创建新列表存储有效file_id
                    valid_file_ids = []
                    for file_id in file_ids:
                        # 使用参数化查询防止SQL注入 [[5]]
                        query = "SELECT id, path FROM file WHERE id = %s"
                        cursor.execute(query, (file_id,))
                        row = cursor.fetchone()  # 假设id唯一，使用fetchone()

                        if not row:  # 未查询到结果时跳过
                            continue

                        current_path = row[1]
                        # 检查路径是否以HTTP开头 [[7]]
                        if current_path.lower().startswith(('http://', 'https://')):
                            continue  # 跳过HTTP路径

                        # 保留有效file_id和路径
                        valid_file_ids.append(file_id)
                        file_dict[int(file_id)] = current_path  # 更新路径
                        file_dict_rev[current_path] = int(file_id)
                    # 更新result中的有效file_ids列表
                    if valid_file_ids:
                        result[entity][1] = valid_file_ids
        except Exception as e:
            logger.error(f"Error occurred: {e}")

        return file_dict

    def get_dir_private_list(self, entity_id, userid):
        """
        查询dir_entity，如果没结果则查询dir_to_entity -> directory，并写入新内容后返回third列表（有事务控制版本）
        """
        # 开始事务
        connection = self.connection
        try:
            # 开始事务
            connection.begin()

            with connection.cursor() as cursor:
                # 先查 dir_entity
                sql1 = """
                       SELECT dir_private
                       FROM dir_entity
                       WHERE entity_id = %s \
                         AND userid = %s \
                       """
                cursor.execute(sql1, (entity_id, userid))
                result = cursor.fetchall()
                if result:
                    # 如果查到结果，提交事务并返回
                    connection.commit()
                    return [row[0] for row in result]

                # 查不到 -> 查 dir_to_entity 获取 second
                sql2 = "SELECT second FROM dir_to_entity WHERE entity_id = %s"
                cursor.execute(sql2, (entity_id,))
                second_result = cursor.fetchone()
                if not second_result:
                    connection.commit()
                    return []  # 如果连 second 都查不到，直接返回空列表

                second = second_result[0]

                # 查 directory 表获取 third
                sql3 = "SELECT third FROM directory WHERE second = %s"
                cursor.execute(sql3, (second,))
                third_results = cursor.fetchall()
                third_list = [row[0] for row in third_results]

                # 批量插入到 dir_entity 表
                insert_sql = """
                             INSERT INTO dir_entity (entity_id, dir_private, dir_sys, userid)
                             VALUES (%s, %s, %s, %s) \
                             """
                for third in third_list:
                    cursor.execute(insert_sql, (entity_id, third, third, userid))

            # 所有操作成功，提交事务
            connection.commit()
            return third_list

        except Exception as e:
            # 发生错误，回滚事务
            connection.rollback()
            logger.error(f"获取目录私有列表失败，事务已回滚：{e}")
            raise

    def get_dir_private_list_simple(self, entity_id, userid):
        """
        查询dir_entity，如果没结果则查询dir_to_entity -> directory，并写入新内容后返回third列表（无事务控制版本）
        """
        try:
            with self.connection.cursor() as cursor:
                # 先查 dir_entity
                sql1 = """
                       SELECT dir_private
                       FROM dir_entity
                       WHERE entity_id = %s \
                         AND userid = %s \
                       """
                cursor.execute(sql1, (entity_id, userid))
                result = cursor.fetchall()
                if result:
                    # 如果查到结果，直接返回
                    return [row[0] for row in result]

                # 查不到 -> 查 dir_to_entity 获取 second
                sql2 = "SELECT second FROM dir_to_entity WHERE entity_id = %s"
                cursor.execute(sql2, (entity_id,))
                second_result = cursor.fetchone()
                if not second_result:
                    return []  # 如果连 second 都查不到，直接返回空列表

                second = second_result[0]

                # 查 directory 表获取 third
                sql3 = "SELECT third FROM directory WHERE second = %s"
                cursor.execute(sql3, (second,))
                third_results = cursor.fetchall()
                third_list = [row[0] for row in third_results]

                # 批量插入到 dir_entity 表
                insert_sql = """
                             INSERT INTO dir_entity (entity_id, dir_private, dir_sys, userid)
                             VALUES (%s, %s, %s, %s) \
                             """
                for third in third_list:
                    cursor.execute(insert_sql, (entity_id, third, third, userid))
            self.connection.commit()
            return third_list

        except Exception as e:
            logger.error(f"获取目录私有列表失败：{e}")
            raise e  # 抛出异常让外层事务处理

    def insert_classification_and_entity_data(self, classification_data, entity_id, file_dict_rev=None, userid=None):
        """
        合并的函数：插入分类结果和目录到实体关联
        参数:
        - classification_data: API返回的分类结果数据
        - entity_id: 实体ID
        - file_dict_rev: 文件字典反向映射（可选，用于分类结果插入）
        - userid: 用户ID（可选，用于分类结果插入）
        """
        # 手动管理Redis分布式锁
        lock = get_distributed_lock(f"classification_entity:{entity_id}", timeout=30)

        try:
            # 获取分布式锁 - 阻塞模式，使用Redis原生BLPOP等待
            logger.info(f"⏳ 线程 {threading.current_thread().name} 正在等待获取分类插入分布式锁: entity_id={entity_id}")
            lock.acquire(blocking=True)  # 使用Redis原生阻塞机制
            logger.info(f"🔒 线程 {threading.current_thread().name} 成功获取分类插入分布式锁: entity_id={entity_id}")

            try:
                data = classification_data['data']
                files = data.get('files', {})
                label_1 = data.get('label_1', '')
                label_2 = data.get('label_2', '')

                # 第一部分：插入目录到实体关联（原 insert_dir_toentity 的功能）
                dir_entity_data = {
                    'second': label_2,
                    'entity_id': entity_id,
                }
                self.insert_data_without_primary('dir_to_entity', dir_entity_data)
                logger.info(f"成功插入目录到实体关联: entity_id={entity_id}, label_2={label_2}")

                # 第二部分：插入分类结果（原 insert_classification_result 的功能）
                results = []
                if file_dict_rev is not None and userid is not None and files:
                    # 为每个文件插入记录
                    for filename, file_category in files.items():
                        # 获取目录ID（无事务版本）
                        directory_id = self.get_directory_simple(label_2, file_category)

                        # 插入到 dir_to_file 表
                        file_data = {
                            'id': directory_id,
                            'fileid': file_dict_rev[f"bb/{filename}"],
                        }
                        self.insert_data_without_primary('dir_to_file', file_data)

                        # 获取私有目录列表（无事务版本）
                        self.get_dir_private_list_simple(entity_id, userid)

                        # 获取新目录ID
                        pri_dir_id = self.get_new_directory(entity_id, userid, file_category)

                        # 插入到 dir_file 表
                        dir_file_data = {
                            'dir_id': pri_dir_id[0],
                            'file_id': file_dict_rev[f"bb/{filename}"],
                        }
                        self.insert_data_without_primary("dir_file", dir_file_data)

                    logger.info(f"成功插入分类结果: entity_id={entity_id}, 文件数量={len(files)}")

                logger.info(f"所有数据插入成功: entity_id={entity_id}")
                logger.info(f"✅ 分类插入操作完成 (锁将释放): entity_id={entity_id}")
                return True

            except Exception as inner_e:
                error_msg = f"分类数据插入失败 - entity_id: {entity_id}, 错误: {str(inner_e)}"
                logger.error(f"❌ {error_msg}")
                # 发生异常时尝试释放锁
                try:
                    lock.release()
                    logger.info(
                        f"🔓 线程 {threading.current_thread().name} 异常释放分类插入分布式锁: entity_id={entity_id}")
                except Exception as release_error:
                    logger.error(f"❌ 异常处理中释放分布式锁失败 - entity_id: {entity_id}, 错误: {str(release_error)}")
                raise  # 重新抛出原始异常

        except Exception as outer_e:
            error_msg = str(outer_e)
            detailed_msg = f"分类插入分布式锁操作异常: {error_msg}"
            logger.error(f"❌ {detailed_msg}")

            # 如果是Redis连接问题，提供更具体的错误信息
            if "Redis" in error_msg or "Connection" in error_msg:
                detailed_msg += " (Redis服务器可能未启动或连接配置有问题)"
                logger.warning(f"💡 Redis服务器可能未启动或连接配置有问题")

            # 确保锁被释放
            try:
                lock.release()
                logger.info(
                    f"🔓 线程 {threading.current_thread().name} 外层异常释放分类插入分布式锁: entity_id={entity_id}")
            except Exception as release_error:
                logger.error(f"❌ 外层异常处理中释放分布式锁失败 - entity_id: {entity_id}, 错误: {str(release_error)}")

            raise  # 重新抛出异常

        finally:
            # 最终释放锁
            try:
                lock.release()
                logger.info(f"🔓 线程 {threading.current_thread().name} 最终释放分类插入分布式锁: entity_id={entity_id}")
            except Exception as release_error:
                logger.error(f"❌ 最终释放分布式锁失败 - entity_id: {entity_id}, 错误: {str(release_error)}")
                # 在finally块中无法通过return停止，但可以抛出异常
                raise Exception(f"最终释放分布式锁失败 - entity_id: {entity_id}, 错误: {str(release_error)}")

    def upload_direct(self, xiaoqi, second_classify):
        # 验证输入数据有效性
        if not second_classify.get('data', {}).get('entity_path'):
            logger.warning(f"Invalid input data: {second_classify}")
            return
        original_path = second_classify['data']['entity_path']

        # 使用Redis分布式锁，确保跨进程/跨请求的互斥访问
        try:
            with get_distributed_lock(xiaoqi, timeout=60) as lock:
                logger.info(f"🔒 获取到分布式锁: {xiaoqi}")

                # 路径转换逻辑
                parts = original_path.strip('\\').split('\\')
                processed_parts = ['KO目录'] + [p for p in parts]
                new_entity_path = ['->'.join(processed_parts)]

                # 数据库操作 - 添加事务支持
                with self.connection.cursor() as cursor:
                    try:
                        # 开始事务
                        self.connection.begin()

                        # 检查目录状态（注意表名保持一致性）
                        query_sql = "SELECT directory FROM xiaoqi_new WHERE xiaoqi_name = %s"
                        cursor.execute(query_sql, (xiaoqi,))
                        result = cursor.fetchone()

                        if result is None or result[0] is None:
                            # 使用JSON序列化替代字符串强转（更规范的存储方式）
                            directory_json = json.dumps(new_entity_path, ensure_ascii=False)
                            update_sql = "UPDATE xiaoqi_new SET directory = %s WHERE xiaoqi_name = %s"

                            cursor.execute(update_sql, (directory_json, xiaoqi))

                            # 检查是否实际更新了数据
                            if cursor.rowcount > 0:
                                logger.info(f"✅ 成功更新目录: {xiaoqi}")
                            else:
                                logger.warning(f"⚠️ 未找到对应的xiaoqi记录: {xiaoqi}")
                        else:
                            logger.info(f"⚠️ 跳过更新: {xiaoqi} 已存在目录")

                        # 提交事务
                        self.connection.commit()

                    except Exception as e:
                        # 发生异常时回滚事务
                        self.connection.rollback()
                        logger.error(f"❌ 数据库操作失败，已回滚: {str(e)}")
                        raise

                logger.info(f"🔓 释放分布式锁: {xiaoqi}")

        except Exception as e:
            logger.error(f"❌ 分布式锁操作或处理异常: {str(e)}")
            raise  # 重新抛出异常让调用方处理

    def get_new_directory(self, entity_id, userid, label_2):
        """
        根据entity_id和userid查询dir_private字段，并以列表形式返回
        """
        with self.connection.cursor() as cursor:
            sql = """
                  SELECT id
                  FROM dir_entity
                  WHERE entity_id = %s \
                    AND userid = %s \
                    AND dir_private = %s \
                  """
            cursor.execute(sql, (entity_id, userid, label_2))
            result = cursor.fetchall()
            return result[0]

    def get_directory(self, label_1, label_2):
        """
        插入或获取目录ID（有事务控制版本）
        """
        try:
            # 检查目录是否已存在
            query = "SELECT id FROM directory WHERE second = %s AND third = %s"
            with self.connection.cursor() as cursor:
                cursor.execute(query, (label_1, label_2))
                result = cursor.fetchone()

                if result:
                    return result[0]  # 返回已存在的目录ID
                else:
                    return -1
        except pymysql.MySQLError as e:
            logger.error(f"目录插入或查询失败: {e}")
            self.connection.rollback()
            return None

    def get_directory_simple(self, label_1, label_2):
        """
        插入或获取目录ID（无事务控制版本）
        """
        try:
            # 检查目录是否已存在
            query = "SELECT id FROM directory WHERE second = %s AND third = %s"
            with self.connection.cursor() as cursor:
                cursor.execute(query, (label_1, label_2))
                result = cursor.fetchone()

                if result:
                    return result[0]  # 返回已存在的目录ID
                else:
                    return -1
        except pymysql.MySQLError as e:
            logger.error(f"目录插入或查询失败: {e}")
            raise e  # 抛出异常让外层事务处理

    def query_dir_by_name_id(self, name_id):
        """
        根据目录名称和ID查询目录
        参数:
            name_id: 目录名称和ID，格式为"名称:ID"
        返回:
            目录信息，格式为{'id': ID, 'name': 名称}，如果查询失败则返回None
        """
        try:
            with self.connection.cursor() as cursor:
                # 执行查询
                sql = "SELECT second  FROM dir_to_entity WHERE entity_id = %s"
                cursor.execute(sql, (name_id,))
                for second in cursor.fetchall():
                    return second[0]
        except pymysql.MySQLError as e:
            logger.error(f"查询失败：{e}")
            return ""

    def query_entities_by_name(self, name, exact_match=False):
        """
        查询实体并格式化成目标结构
        参数:
            name: 要查询的名称
            exact_match: 是否精确匹配，False为模糊匹配（默认），True为精确匹配
        返回:
            {
                "汪萌1": ["汪萌", "教授", ...],
                "汪萌2": ["重庆大学"],
                ...
            }
        """
        entity_list = {}
        entity_dict = {}
        try:
            self.connect()
            with self.connection.cursor() as cursor:
                # 执行模糊查询 - 支持字符+任意字符的匹配
                sql = """SELECT xiaoqi_name, key_words, xiaoqi_id
                         FROM xiaoqi_new
                         WHERE xiaoqi_name LIKE %s"""
                # 添加调试日志
                search_pattern = f"{name}%"
                logger.debug(f"🔍 查询SQL: {sql}")
                logger.debug(f"🔍 查询参数: {search_pattern}")
                cursor.execute(sql, (search_pattern,))  # 真正添加%通配符，匹配以name开头的任意字符

                # 处理查询结果
                results = cursor.fetchall()
                logger.debug(f"🔍 查询结果数量: {len(results)}")
                for xiaoqi_name, keywords, xiaoqi_id in results:
                    logger.debug(f"🔍 找到记录: xiaoqi_name='{xiaoqi_name}', xiaoqi_id={xiaoqi_id}")

                    # 添加JSON解析的错误处理
                    if keywords is None or keywords.strip() == "":
                        keyword_list = [xiaoqi_name]
                    else:
                        try:
                            keywords = keywords.replace("'", '"').replace("'", '"').replace("'", '"')
                            keyword_list = json.loads(keywords)
                        except json.JSONDecodeError as e:
                            logger.warning(
                                f"⚠️ JSON解析失败，使用默认值: xiaoqi_name='{xiaoqi_name}', keywords='{keywords}', 错误: {e}")
                            keyword_list = [xiaoqi_name]

                    entity_list[xiaoqi_name] = keyword_list
                    entity_dict[xiaoqi_name] = xiaoqi_id
        except pymysql.MySQLError as e:
            logger.error(f"查询失败：{e}")
            return {}
        finally:
            self.close()
        return entity_list, entity_dict

    def connect(self):
        """
        建立数据库连接
        """
        try:
            self.connection = pymysql.connect(**self.config)
            # 确保开启自动提交模式
            self.connection.autocommit(True)
            logger.info(f"✅ 数据库连接成功，已开启autocommit")
        except pymysql.MySQLError as e:
            logger.error(f"数据库连接失败：{e}")
            raise

    def insert_xiaoqi_new(self, table_name, data, primary_key='xiaoqi_name'):
        # 获取主键值
        primary_key_value = data.get(primary_key)

        with self.connection.cursor() as cursor:
            # 检查记录是否已存在
            check_query = f"SELECT {primary_key}, xiaoqi_id, key_words FROM {table_name} WHERE {primary_key} = %s FOR UPDATE"
            cursor.execute(check_query, (primary_key_value,))
            existing_record = cursor.fetchone()

            if existing_record:
                # 记录存在
                existing_name, existing_id, existing_keywords = existing_record

                if existing_keywords is not None and str(existing_keywords).strip() != "":
                    # 记录存在且 key_words 不为空，返回现有ID
                    logger.info(f"主键 {primary_key_value} 已存在且key_words不为空，返回现有ID: {existing_id}")
                    return existing_id
                else:
                    # 记录存在但 key_words 为空，执行更新操作
                    logger.info(f"主键 {primary_key_value} 已存在但key_words为空，执行更新操作")

                    # 构建更新语句，只更新非主键字段
                    update_fields = []
                    update_values = []
                    for key, value in data.items():
                        if key != primary_key:  # 排除主键字段
                            update_fields.append(f"{key} = %s")
                            update_values.append(value)

                    if update_fields:
                        update_query = f"UPDATE {table_name} SET {', '.join(update_fields)} WHERE {primary_key} = %s"
                        update_values.append(primary_key_value)
                        cursor.execute(update_query, tuple(update_values))
                        self.connection.commit()
                        logger.info(f"✅ 成功更新xiaoqi: {primary_key_value}")
                        logger.info(f"🔥 insert_xiaoqi_new 更新提交成功: {primary_key_value}")

                    return existing_id
            else:
                # 记录不存在，插入新数据
                logger.info(f"主键 {primary_key_value} 不存在，执行插入操作")
                columns = ", ".join(data.keys())
                placeholders = ", ".join(["%s"] * len(data))
                insert_query = f"INSERT INTO {table_name} ({columns}) VALUES ({placeholders})"
                cursor.execute(insert_query, tuple(data.values()))
                self.connection.commit()
                logger.info(f"✅ 成功插入xiaoqi: {primary_key_value}")
                logger.info(f"🔥 insert_xiaoqi_new 插入提交成功: {primary_key_value}")
                return cursor.lastrowid

    def insert_data(self, table_name, data):
        try:
            # 先检查主键是否存在
            primary_key = list(data.keys())[0]  # 假设主键在第一个位置
            primary_key_value = data[primary_key]

            # 生成检查主键是否存在的 SQL 查询
            check_query = f"SELECT COUNT(*) FROM {table_name} WHERE {primary_key} = %s"
            with self.connection.cursor() as cursor:
                cursor.execute(check_query, (primary_key_value,))
                result = cursor.fetchone()

                if result[0] > 0:
                    return  # 主键已存在，跳过插入操作

            # 生成插入 SQL 语句
            columns = ", ".join(data.keys())
            placeholders = ", ".join(["%s"] * len(data))
            insert_query = f"INSERT INTO {table_name} ({columns}) VALUES ({placeholders})"

            # 执行插入操作
            with self.connection.cursor() as cursor:
                cursor.execute(insert_query, tuple(data.values()))
        except pymysql.MySQLError as e:
            logger.error(f"插入数据失败：{e}")

    def insert_relation(self, table_name, data):
        try:
            # 先检查主键是否存在
            primary_key = list(data.keys())[0]  # 假设主键在第一个位置
            last_key = list(data.keys())[-1]
            primary_key_value = data[primary_key]
            last_key_value = data[last_key]

            # 生成检查主键是否存在的 SQL 查询
            check_query = f"SELECT COUNT(*) FROM {table_name} WHERE {primary_key} = %s AND {last_key} = %s "
            with self.connection.cursor() as cursor:
                cursor.execute(check_query, (primary_key_value, last_key_value,))
                result = cursor.fetchone()

                if result[0] > 0:
                    return  # 主键已存在，跳过插入操作

            # 生成插入 SQL 语句
            columns = ", ".join(data.keys())
            placeholders = ", ".join(["%s"] * len(data))
            insert_query = f"INSERT INTO {table_name} ({columns}) VALUES ({placeholders})"

            # 执行插入操作
            with self.connection.cursor() as cursor:
                cursor.execute(insert_query, tuple(data.values()))
        except pymysql.MySQLError as e:
            logger.error(f"插入数据失败：{e}")

    def close(self):
        """
        关闭数据库连接
        """
        if self.connection:
            self.connection.close()
            logger.info("数据库连接已关闭！")


class Bucket:

    def __init__(self, minio_address, minio_admin, minio_password):
        # 通过ip 账号 密码 连接minio server
        # Http连接 将secure设置为False
        self.minioClient = Minio(endpoint=minio_address,
                                 access_key=minio_admin,
                                 secret_key=minio_password,
                                 secure=False)

    def create_one_bucket(self, bucket_name):
        # 创建桶(调用make_bucket api来创建一个桶)
        """
        桶命名规则：小写字母，句点，连字符和数字 允许使用 长度至少3个字符
        使用大写字母、下划线等会报错
        """
        try:
            # bucket_exists：检查桶是否存在
            if not self.minioClient.bucket_exists(bucket_name=bucket_name):
                self.minioClient.make_bucket(bucket_name=bucket_name)
        except InvalidResponseError as err:
            logger.error(f"创建桶失败: {err}")

    def remove_one_bucket(self, bucket_name):
        # 删除桶(调用remove_bucket api来创建一个存储桶)
        try:
            if self.minioClient.bucket_exists(bucket_name=bucket_name):
                self.minioClient.remove_bucket(bucket_name)
            else:
                logger.warning("该存储桶不存在")
        except InvalidResponseError as err:
            logger.error(f"删除桶失败: {err}")

    def upload_stream_tobucket(self, bucket_name, upload_file):
        object_name = upload_file.name  # 或者自定义带前缀的路径
        try:
            minio_client.put_object(
                bucket_name=bucket_name,
                object_name=object_name,
                data=upload_file,  # upload_file 本身支持 .read()
                length=upload_file.size,  # 请求对象的总大小
                content_type=upload_file.content_type  # 可选，推荐提供
            )
        except S3Error as err:
            return {"error": f"上传失败：{err}"}

    def upload_file_to_bucket(self, bucket_name, file_name, file_path):
        """
        将文件上传到bucket
        :param bucket_name: minio桶名称
        :param file_name: 存放到minio桶中的文件名字(相当于对文件进行了重命名，可以与原文件名不同)
                            file_name处可以创建新的目录(文件夹) 例如 /example/file_name
                            相当于在该桶中新建了一个example文件夹 并把文件放在其中
        :param file_path: 本地文件的路径
        """
        # 桶是否存在 不存在则新建
        check_bucket = self.minioClient.bucket_exists(bucket_name)
        if not check_bucket:
            self.minioClient.make_bucket(bucket_name)

        try:
            self.minioClient.fput_object(bucket_name=bucket_name,
                                         object_name=file_name,
                                         file_path=file_path)
        except FileNotFoundError as err:
            logger.error(f'upload_failed: {str(err)}')
        except S3Error as err:
            logger.error(f"upload_failed: {err}")

    def download_file_from_bucket(self, bucket_name, minio_file_path, download_file_path):
        """
        从bucket下载文件
        :param bucket_name: minio桶名称
        :param minio_file_path: 存放在minio桶中文件名字
                            file_name处可以包含目录(文件夹) 例如 /example/file_name
        :param download_file_path: 文件获取后存放的路径
        """
        # 桶是否存在
        check_bucket = self.minioClient.bucket_exists(bucket_name)
        if check_bucket:
            try:
                self.minioClient.fget_object(bucket_name=bucket_name,
                                             object_name=minio_file_path,
                                             file_path=download_file_path)
                return 1
            except FileNotFoundError as err:
                logger.error(f'download_failed: {str(err)}')
                return 0
            except S3Error as err:
                logger.error(f"download_failed: {err}")

    def remove_object(self, bucket_name, object_name):
        """
        从bucket删除文件
        :param bucket_name: minio桶名称
        :param object_name: 存放在minio桶中的文件名字
                            object_name处可以包含目录(文件夹) 例如 /example/file_name
        """
        # 桶是否存在
        check_bucket = self.minioClient.bucket_exists(bucket_name)
        if check_bucket:
            try:
                self.minioClient.remove_object(bucket_name=bucket_name,
                                               object_name=object_name)
            except FileNotFoundError as err:
                logger.error(f'delete_failed: {str(err)}')
            except S3Error as err:
                logger.error(f"delete_failed: {err}")

    # 获取所有的桶
    def get_all_bucket(self):
        buckets = self.minioClient.list_buckets()
        ret = []
        for _ in buckets:
            ret.append(_.name)
        return ret

    # 获取一个桶中的所有一级目录和文件
    def get_list_objects_from_bucket(self, bucket_name):
        # 桶是否存在
        check_bucket = self.minioClient.bucket_exists(bucket_name)
        if check_bucket:
            # 获取到该桶中的所有目录和文件
            objects = self.minioClient.list_objects(bucket_name=bucket_name)
            ret = []
            for _ in objects:
                ret.append(_.object_name)
            return ret

    # 获取桶里某个目录下的所有目录和文件
    def get_list_objects_from_bucket_dir(self, bucket_name, dir_name):
        # 桶是否存在
        check_bucket = self.minioClient.bucket_exists(bucket_name)
        if check_bucket:
            objects = self.minioClient.list_objects(bucket_name=bucket_name,
                                                    prefix=dir_name,
                                                    recursive=True)
            ret = []
            for obj in objects:
                object_name = obj.object_name
                # 获取对象的内容
                content = self.minioClient.get_object(bucket_name=bucket_name,
                                                      object_name=object_name)
                ret.append(content.data.decode())
            return ret


def get_sha1_hash(file_name):
    shal_hash = hashlib.sha1(file_name.encode()).hexdigest()
    return shal_hash


def read_txt_file(file_path):
    with open(file_path, 'r', encoding="utf-8") as file:
        file_contents = file.read()
    return file_contents


class HtmlParse(HTMLParser):
    """
    基于HTMLParser的HTML解析器
    提取指定标签的内容和属性
    """

    def __init__(self, flag=2):
        super().__init__()
        self.flag = flag
        self.main_content = []
        self.image_content = []

        # 状态变量
        self.in_p_tag = False
        self.in_span_tag = False
        self.in_a_tag = False
        self.current_span_class = None
        self.current_text = ""
        self.span_inner_text = ""

        # 临时存储
        self.temp_meta_attrs = {}

    def handle_starttag(self, tag, attrs):
        """处理开始标签"""
        attrs_dict = dict(attrs)

        if tag == 'meta':
            self.temp_meta_attrs = attrs_dict

        elif tag == 'p' and self.flag == 1:  # 维基百科段落
            self.in_p_tag = True
            self.current_text = ""

        elif tag == 'span' and self.flag == 2:  # 百度百科span
            if 'class' in attrs_dict and 'text_bypwF' in attrs_dict['class']:
                self.in_span_tag = True
                self.current_span_class = attrs_dict['class']
                self.span_inner_text = ""

        elif tag == 'a' and self.in_span_tag:
            self.in_a_tag = True

    def handle_endtag(self, tag):
        """处理结束标签"""
        if tag == 'meta':
            self._process_meta_tag()
            self.temp_meta_attrs = {}

        elif tag == 'p' and self.in_p_tag:
            self.in_p_tag = False
            if self.current_text.strip():
                self.main_content.append(self.current_text.strip())
            self.current_text = ""

        elif tag == 'span' and self.in_span_tag:
            self.in_span_tag = False
            if self.span_inner_text.strip():
                self.main_content.append(self.span_inner_text.strip())
            self.span_inner_text = ""

        elif tag == 'a' and self.in_a_tag:
            self.in_a_tag = False

    def handle_data(self, data):
        """处理文本数据"""
        if self.in_p_tag:
            self.current_text += data

        elif self.in_span_tag:
            self.span_inner_text += data

    def _process_meta_tag(self):
        """处理meta标签内容"""
        if not self.temp_meta_attrs:
            return

        # 处理描述内容
        if 'name' in self.temp_meta_attrs:
            name = self.temp_meta_attrs['name'].lower()
            if name == 'description' and 'content' in self.temp_meta_attrs:
                content = html.unescape(self.temp_meta_attrs['content'])
                self.main_content.append(content)
            elif name == 'image' and 'content' in self.temp_meta_attrs:
                content = self.temp_meta_attrs['content']
                self.image_content.append(content)

        elif 'property' in self.temp_meta_attrs:
            property_val = self.temp_meta_attrs['property'].lower()
            if property_val == 'og:description' and 'content' in self.temp_meta_attrs:
                content = html.unescape(self.temp_meta_attrs['content'])
                self.main_content.append(content)
            elif property_val == 'og:image' and 'content' in self.temp_meta_attrs:
                content = self.temp_meta_attrs['content']
                self.image_content.append(content)

    def get_results(self):
        """获取解析结果"""
        return self.main_content, self.image_content


def Dir_html_word(html_file_path, flag):
    """
    使用HtmlParse解析HTML文件

    Args:
        html_file_path (str): HTML文件路径
        flag (int): 解析模式 1-维基百科 2-百度百科

    Returns:
        dict: 统一的解析结果格式
        {
            "success": bool,          # 解析是否成功
            "code": int,              # 状态码 200-成功，400-客户端错误，500-服务器错误
            "message": str,           # 状态消息
            "data": {                 # 解析结果数据
                "content": list,      # 解析出的主要内容列表
                "image_content": list,# 图片内容（如果有）
                "content_count": int, # 内容条数
                "file_info": dict     # 文件信息
            },
            "error": str              # 错误详情（如果有）
        }
    """
    try:
        html_content = None
        file_basename = os.path.basename(html_file_path)

        # 读取HTML文件并检测编码
        with open(html_file_path, 'rb') as f:
            raw_data = f.read()
            detected = chardet.detect(raw_data)
            encoding = detected['encoding'] if detected['confidence'] > 0.7 else 'utf-8'

            try:
                html_content = raw_data.decode(encoding, errors='replace')
            except (LookupError, UnicodeDecodeError):
                # 如果检测的编码无效，尝试常见中文编码
                try:
                    html_content = raw_data.decode('gbk', errors='replace')
                except UnicodeDecodeError:
                    html_content = raw_data.decode('utf-8', errors='replace')

        # 检测反爬虫和异常内容
        if is_anti_crawler_or_error_page(html_content):
            logger.warning(f"⚠️ 检测到反爬虫或错误页面: {html_file_path}")
            return {
                "success": False,
                "code": 400,
                "message": "检测到反爬虫或错误页面",
                "data": {
                    "content": [],
                    "image_content": [],
                    "content_count": 0,
                    "file_info": {
                        "file_name": file_basename,
                        "file_path": html_file_path,
                        "encoding": encoding,
                        "parse_mode": flag
                    }
                },
                "error": f"文件 {file_basename} 包含无效内容或被访问限制"
            }

        # 使用HtmlParse解析HTML内容
        parser = HtmlParse(flag=flag)
        parser.feed(html_content)

        main_content, image_content = parser.get_results()

        # 如果使用百度百科模式但没有找到span内容，尝试提取所有文本
        if flag == 2 and not main_content:
            logger.warning(f"未找到百度百科特定内容，尝试提取通用文本内容: {html_file_path}")
            # 创建通用解析器提取所有可见文本
            generic_parser = GenericHtmlParser()
            generic_parser.feed(html_content)
            main_content = generic_parser.get_text_content()

        # 再次检查解析后的内容是否有效
        if main_content:
            filtered_content = filter_invalid_content(main_content)
            if not filtered_content:
                logger.warning(f"⚠️ 解析后内容被过滤为空: {html_file_path}")
                return {
                    "success": False,
                    "code": 400,
                    "message": "解析后内容无效",
                    "data": {
                        "content": [],
                        "image_content": image_content or [],
                        "content_count": 0,
                        "file_info": {
                            "file_name": file_basename,
                            "file_path": html_file_path,
                            "encoding": encoding,
                            "parse_mode": flag
                        }
                    },
                    "error": f"文件 {file_basename} 解析后内容无效"
                }
            main_content = filtered_content

        parser.close()

        # 确保返回的内容是列表格式
        if not isinstance(main_content, list):
            main_content = [main_content] if main_content else []
        if not isinstance(image_content, list):
            image_content = [image_content] if image_content else []

        logger.info(f"✅ HTML解析完成: {html_file_path}, 提取内容数量: {len(main_content)}")

        return {
            "success": True,
            "code": 200,
            "message": "HTML解析成功",
            "data": {
                "content": main_content,
                "image_content": image_content,
                "content_count": len(main_content),
                "file_info": {
                    "file_name": file_basename,
                    "file_path": html_file_path,
                    "encoding": encoding,
                    "parse_mode": flag
                }
            },
            "error": None
        }

    except FileNotFoundError:
        logger.error(f"❌ 文件未找到: {html_file_path}")
        return {
            "success": False,
            "code": 404,
            "message": "文件未找到",
            "data": {
                "content": [],
                "image_content": [],
                "content_count": 0,
                "file_info": {
                    "file_name": os.path.basename(html_file_path),
                    "file_path": html_file_path,
                    "encoding": None,
                    "parse_mode": flag
                }
            },
            "error": f"文件未找到: {os.path.basename(html_file_path)}"
        }
    except IOError as e:
        logger.error(f"❌ 文件读取错误 {html_file_path}: {e}")
        return {
            "success": False,
            "code": 500,
            "message": "文件读取错误",
            "data": {
                "content": [],
                "image_content": [],
                "content_count": 0,
                "file_info": {
                    "file_name": os.path.basename(html_file_path),
                    "file_path": html_file_path,
                    "encoding": None,
                    "parse_mode": flag
                }
            },
            "error": f"文件读取错误: {str(e)}"
        }
    except Exception as e:
        logger.error(f"❌ HTML解析错误 {html_file_path}: {e}")
        return {
            "success": False,
            "code": 500,
            "message": "HTML解析失败",
            "data": {
                "content": [],
                "image_content": [],
                "content_count": 0,
                "file_info": {
                    "file_name": os.path.basename(html_file_path),
                    "file_path": html_file_path,
                    "encoding": None,
                    "parse_mode": flag
                }
            },
            "error": f"HTML解析失败: {str(e)}"
        }


def is_anti_crawler_or_error_page(html_content):
    """
    检测是否为反爬虫页面或错误页面

    Args:
        html_content (str): HTML内容

    Returns:
        bool: True表示是反爬虫或错误页面
    """
    if not html_content or len(html_content.strip()) < 50:
        return True

    # 常见的反爬虫和错误页面特征
    anti_crawler_indicators = [
        # 知乎相关
        "您当前请求存在异常，暂时限制本次访问",
        "知乎小管家反馈",
        "摇一摇或登录后私信",
        "7459a2388cb8a6eed61551b9eec0924",

        # 百度相关
        "百度安全验证",
        "请输入验证码",
        "访问过于频繁",
        "系统检测到您的网络环境存在异常",

        # 通用反爬虫
        "访问被拒绝",
        "Access Denied",
        "403 Forbidden",
        "Please verify you are human",
        "robot or spider",
        "verification required",
        "captcha",
        "频繁访问",
        "访问限制",
        "暂时无法访问",
        "请稍后再试",
        "Rate limit exceeded",
        "Too many requests",

        # 错误页面
        "404 Not Found",
        "页面不存在",
        "找不到页面",
        "服务器错误",
        "Server Error",
        "Internal Server Error",
        "网络错误",
        "连接超时",
        "请求超时",

        # 重定向或加载页面
        "正在跳转",
        "页面跳转中",
        "Loading...",
        "请稍候",
        "加载中",

        # 登录页面
        "请登录",
        "Please login",
        "Sign in required",
        "登录后查看",
        "需要登录"
    ]

    # 检查内容长度，过短的内容通常无效
    if len(html_content.strip()) < 100:
        return True

    # 检查是否包含反爬虫特征
    html_lower = html_content.lower()
    for indicator in anti_crawler_indicators:
        if indicator.lower() in html_lower:
            logger.warning(f"🚨 检测到反爬虫指标: {indicator}")
            return True

    # 检查是否主要由脚本或样式组成（无实际内容）
    try:
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(html_content, 'html.parser')

        # 移除脚本和样式标签
        for script in soup(["script", "style"]):
            script.decompose()

        # 获取纯文本内容
        text_content = soup.get_text(strip=True)

        # 如果纯文本内容过少，认为是无效页面
        if len(text_content) < 50:
            logger.warning(f"🚨 页面纯文本内容过少: {len(text_content)} 字符")
            return True

        # 检查是否大部分是重复字符或无意义内容
        if is_repetitive_or_meaningless(text_content):
            logger.warning(f"🚨 检测到重复或无意义内容")
            return True
    except Exception as e:
        logger.warning(f"⚠️ 内容检测过程中出现异常: {e}")
        return True

    return False


def is_repetitive_or_meaningless(text):
    """
    检测文本是否为重复或无意义内容

    Args:
        text (str): 文本内容

    Returns:
        bool: True表示是重复或无意义内容
    """
    if not text or len(text.strip()) < 20:
        return True

    # 检查是否大部分是重复字符
    char_counts = {}
    for char in text:
        char_counts[char] = char_counts.get(char, 0) + 1

    # 如果某个字符占比超过30%，认为是重复内容
    total_chars = len(text)
    for char, count in char_counts.items():
        if count / total_chars > 0.3 and char not in [' ', '\n', '\t']:
            return True

    # 检查是否包含大量数字和特殊字符（可能是加密或混淆内容）
    import re
    special_chars = len(re.findall(r'[^a-zA-Z0-9\u4e00-\u9fff\s]', text))
    if special_chars / total_chars > 0.4:
        return True

    # 检查是否有基本的中文或英文内容
    chinese_chars = len(re.findall(r'[\u4e00-\u9fff]', text))
    english_chars = len(re.findall(r'[a-zA-Z]', text))

    if chinese_chars + english_chars < total_chars * 0.3:
        return True

    return False


def filter_invalid_content(content_list):
    """
    过滤无效内容

    Args:
        content_list (list): 内容列表

    Returns:
        list: 过滤后的内容列表
    """
    if not content_list:
        return []

    filtered = []
    for content in content_list:
        if not content or not isinstance(content, str):
            continue

        content = content.strip()
        if len(content) < 10:  # 过短的内容
            continue

        # 检查是否为纯数字、纯符号或重复字符
        if content.isdigit() or len(set(content)) <= 3:
            continue

        # 检查是否包含有意义的文字
        import re
        meaningful_chars = len(re.findall(r'[\u4e00-\u9fff]', content))  # 中文字符
        meaningful_chars += len(re.findall(r'[a-zA-Z]', content))  # 英文字符

        if meaningful_chars < len(content) * 0.3:  # 有意义字符少于30%
            continue

        # 检查是否为常见的无效内容
        invalid_patterns = [
            r'^[\d\s\-_=+*#@$%^&()[\]{}<>|\\/:;,.?!~`\'\"]+$',  # 纯符号
            r'^(.*?)\1{3,}',  # 重复模式
            r'^\s*$',  # 空白
        ]

        is_invalid = False
        for pattern in invalid_patterns:
            if re.match(pattern, content):
                is_invalid = True
                break

        if not is_invalid:
            filtered.append(content)

    return filtered


def extract_text_from_docx(docx_path):
    try:
        # 打开docx文件
        doc = Document(docx_path)

        # 提取所有段落文本
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)

        # 提取表格中的文本
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)

        # 合并所有文本，用换行符分隔
        return '\n'.join(full_text)

    except Exception as e:
        logger.error(f"Error processing {docx_path}: {str(e)}")
        return None


def extract_text_with_markitdown(file_path, file_name_original):
    """
    使用Microsoft MarkItDown库统一解析各种文档格式

    Args:
        file_path (str): 文件路径
        file_name_original (str): 原始文件名（用作备用内容）

    Returns:
        str: 解析后的文本内容，如果解析失败则返回文件名
    """
    try:
        # 初始化MarkItDown
        md = MarkItDown()

        logger.info(f"🔍 开始使用MarkItDown解析文件: {file_path}")

        # 使用MarkItDown解析文件
        result = md.convert(file_path)

        if result and hasattr(result, 'text_content') and result.text_content:
            # 获取解析后的文本内容
            text_content = result.text_content.strip()[:1000]

            if text_content:
                logger.info(f"✅ MarkItDown解析成功: {file_path}, 内容长度: {len(text_content)} 字符")
                return text_content
            else:
                logger.warning(f"⚠️ MarkItDown解析后内容为空: {file_path}")
                return file_name_original

        # 如果result没有text_content属性，尝试直接使用result
        elif result and hasattr(result, 'content'):
            text_content = result.content.strip()[:1000]
            if text_content:
                logger.info(f"✅ MarkItDown解析成功(使用content): {file_path}, 内容长度: {len(text_content)} 字符")
                return text_content
            else:
                logger.warning(f"⚠️ MarkItDown解析后content为空: {file_path}")
                return file_name_original

        # 如果result是字符串
        elif isinstance(result, str) and result.strip():
            logger.info(f"✅ MarkItDown解析成功(字符串结果): {file_path}, 内容长度: {len(result)} 字符")
            return result.strip()

        # 处理DocumentConverterR对象
        elif result and hasattr(result, '__class__') and 'DocumentConverterR' in str(type(result)):
            logger.info(f"🔍 检测到DocumentConverterR对象，尝试提取内容: {file_path}")
            try:
                # 尝试获取文本内容
                if hasattr(result, 'text_content'):
                    text_content = result.text_content.strip()[:1000] if result.text_content else ""
                elif hasattr(result, 'content'):
                    text_content = result.content.strip()[:1000] if result.content else ""
                elif hasattr(result, 'text'):
                    text_content = result.text.strip()[:1000] if result.text else ""
                elif hasattr(result, 'html'):
                    text_content = result.html.strip()[:1000] if result.html else ""
                else:
                    # 尝试将对象转换为字符串
                    text_content = str(result).strip()[:1000]

                if text_content and len(text_content) > 10:
                    logger.info(f"✅ DocumentConverterR对象解析成功: {file_path}, 内容长度: {len(text_content)} 字符")
                    return text_content
                else:
                    logger.warning(f"⚠️ DocumentConverterR对象内容为空或过短: {file_path}")
                    return file_name_original

            except Exception as e:
                logger.error(f"❌ 处理DocumentConverterR对象时出错: {file_path}, 错误: {str(e)}")
                return file_name_original

        else:
            logger.warning(f"⚠️ MarkItDown返回了未预期的结果格式: {file_path}, 结果类型: {type(result)}")
            # 尝试将结果转换为字符串作为最后的备用方案
            try:
                if result:
                    text_content = str(result).strip()[:1000]
                    if text_content and len(text_content) > 10:
                        logger.info(f"✅ 将结果转换为字符串成功: {file_path}, 内容长度: {len(text_content)} 字符")
                        return text_content
            except Exception as e:
                logger.error(f"❌ 转换结果为字符串时出错: {file_path}, 错误: {str(e)}")

            return file_name_original

    except FileNotFoundError:
        logger.error(f"❌ 文件未找到: {file_path}")
        return file_name_original
    except ImportError as import_e:
        logger.error(f"❌ MarkItDown库导入失败: {str(import_e)}")
        logger.error("💡 请安装markitdown库: pip install markitdown")
        return file_name_original
    except Exception as e:
        error_msg = str(e)
        logger.warning(f"⚠️ MarkItDown解析失败: {file_path}, 错误: {error_msg}")

        # 对于某些已知的错误类型，提供更具体的信息
        if "unsupported" in error_msg.lower():
            logger.info(f"📄 文件格式不受支持，使用文件名作为内容: {file_path}")
        elif "corrupted" in error_msg.lower() or "damaged" in error_msg.lower():
            logger.warning(f"📄 文件可能已损坏: {file_path}")
        elif "permission" in error_msg.lower():
            logger.error(f"📄 文件权限问题: {file_path}")

        # 返回文件名作为备用内容
        return file_name_original


def extract_text_with_markitdown_safe(file_path, file_name_original, doc_type):
    """
    安全的MarkItDown文档解析，带有备用方案

    Args:
        file_path (str): 文件路径
        file_name_original (str): 原始文件名
        doc_type (str): 文档类型

    Returns:
        str: 解析后的文本内容
    """
    # 首先尝试使用MarkItDown
    result = extract_text_with_markitdown(file_path, file_name_original)

    # # 如果MarkItDown解析失败（返回了文件名），并且是特定格式，尝试备用方案
    # if result == file_name_original and doc_type:
    #     logger.info(f"🔄 MarkItDown解析失败，尝试备用解析方案: {doc_type}")

    # try:
    #     if doc_type == "pdf":
    #         # PDF备用方案：使用PyMuPDF
    #         backup_result = extract_text_with_pymupdf(file_path)
    #         if backup_result and backup_result.strip():
    #             logger.info(f"✅ PDF备用方案成功: {file_path}")
    #             return backup_result

    #     elif doc_type == "docx":
    #         # DOCX备用方案：使用python-docx
    #         backup_result = extract_text_from_docx(file_path)
    #         if backup_result and backup_result.strip():
    #             logger.info(f"✅ DOCX备用方案成功: {file_path}")
    #             return backup_result

    #     elif doc_type == "html":
    #         # HTML备用方案：使用现有的HTML解析器
    #         html_result = Dir_html_word(file_path, 2)
    #         if html_result["success"] and html_result["data"]["content"]:
    #             html_content_list = html_result["data"]["content"]
    #             backup_result = '\n'.join(html_content_list)
    #             logger.info(f"✅ HTML备用方案成功: {file_path}")
    #             return backup_result

    #         elif doc_type == "txt":
    #             # TXT备用方案：直接读取文件
    #             backup_result = read_txt_file(file_path)
    #             if backup_result and backup_result.strip():
    #                 logger.info(f"✅ TXT备用方案成功: {file_path}")
    #                 return backup_result

    #     except Exception as backup_e:
    #         logger.warning(f"⚠️ 备用解析方案也失败: {file_path}, 错误: {str(backup_e)}")

    # 如果所有方案都失败，返回原始结果（通常是文件名）
    return result


def generate_entity_json(result, file_text_list, file_dict, info_list, customize_content_list, only_name,
                         file_id_list=None):
    """
    生成实体JSON数据，新格式包含文件内容字典

    Args:
        result: 实体结果字典
        file_text_list: 文件内容列表（按file_id_list顺序排列）
        file_dict: 文件ID到路径的映射
        info_list: 实体信息列表
        customize_content_list: 自定义内容列表
        only_name: 是否只使用实体名（保持兼容性，实际已不使用）
        file_id_list: 文件ID列表（与file_text_list对应）

    Returns:
        list: 格式化的实体数据列表，新格式包含filetext_dict
    """
    output_list = []
    # 确保 info_list 和 customize_content_list 的长度与 result 匹配
    result_len = len(result)
    if len(info_list) < result_len:
        info_list = info_list + [''] * (result_len - len(info_list))

    if len(customize_content_list) < result_len:
        customize_content_list = customize_content_list + [''] * (result_len - len(customize_content_list))

    # 创建file_id到content的映射，确保正确的对应关系
    file_id_to_content = {}
    if file_id_list and len(file_text_list) == len(file_id_list):
        # 使用传入的file_id_list建立正确的映射
        for fid, content in zip(file_id_list, file_text_list):
            file_id_to_content[fid] = content
        logger.info(f"✅ 建立file_id到content映射: {len(file_id_to_content)}个文件")
        # 调试输出映射关系（简化版）
        for idx, (fid, content) in enumerate(zip(file_id_list[:3], file_text_list[:3])):  # 只输出前3个
            content_preview = str(content)[:50] + "..." if len(str(content)) > 50 else str(content)
            logger.debug(f"映射 [{idx}]: file_id={fid} -> content='{content_preview}'")
    else:
        logger.warning(
            f"⚠️ file_id_list参数缺失或长度不匹配，file_text_list长度: {len(file_text_list)}, file_id_list长度: {len(file_id_list) if file_id_list else 0}")

    for entity, info, customize_content in zip(result, info_list, customize_content_list):
        file_ids = result[entity][1]
        minio_paths = [file_dict[int(file_id)] for file_id in file_ids]
        keywords = result[entity][0]
        merged_keywords = list({k: None for k in keywords}.keys())

        # 构建filetext_dict：文件名到文件内容的映射
        filetext_dict = {}

        # 遍历当前实体关联的文件
        for file_id in file_ids:
            try:
                file_path = file_dict[int(file_id)]
                file_name = os.path.basename(file_path)  # 提取文件名

                # 从映射中获取对应的文件内容
                if file_id in file_id_to_content:
                    file_content = file_id_to_content[file_id]
                    logger.debug(f"✅ 通过映射找到文件内容: {file_name}")
                elif int(file_id) in file_id_to_content:
                    file_content = file_id_to_content[int(file_id)]
                    logger.debug(f"✅ 通过映射找到文件内容: {file_name}")
                else:
                    # 如果没有映射，设为空字符串并记录警告
                    file_content = ""
                    logger.warning(f"⚠️ 未找到文件内容映射: file_id={file_id}, file_name={file_name}")

                # 确保内容为字符串类型
                if file_content is None:
                    file_content = ""
                elif not isinstance(file_content, str):
                    file_content = str(file_content)

                # 记录空内容警告
                if not file_content.strip():
                    logger.warning(f"⚠️ 文件内容为空: {file_name}")

                filetext_dict[file_name] = file_content

            except Exception as e:
                logger.error(f"❌ 处理文件失败 file_id={file_id}, 错误: {str(e)}")
                # 即使失败也要保持文件名在字典中
                try:
                    file_path = file_dict[int(file_id)]
                    file_name = os.path.basename(file_path)
                    filetext_dict[file_name] = ""
                except:
                    filetext_dict[f"unknown_file_{file_id}"] = ""

        # 构建新格式的输出
        output = {
            "filetext_dict": filetext_dict,
            "entity": entity,
            "entity_with_keyword": f"{entity}{len(merged_keywords)}" if merged_keywords else entity,  # 新格式：实体名+关键词数量
            "info": info if info is not None else "",
            "customize_content": customize_content
        }
        output_list.append(output)

    return output_list


def updata_to_mysql_new(result, is_xiaoqi):
    db = MySQLDatabase(
        host="114.213.234.179",
        user="koroot",
        password="DMiC-4092",
        database="db_hp"
    )
    entity_file = {}

    try:
        db.connect()  # 确保建立数据库连接
        for key, value in result.items():
            if is_xiaoqi:
                data_to_insert = {
                    "key_words": str(value[0]),
                    "xiaoqi_name": str(key)
                }
                # 插入 xiaoqi_new 表并获取 ID
                xiaoqi_id = db.insert_xiaoqi_new("xiaoqi_new", data_to_insert)
                entity_file[key] = xiaoqi_id

            # 批量处理 xiaoqi_to_file 关联
            file_ids = [int(i) for i in value[1]]
            if file_ids:
                with db.connection.cursor() as cursor:
                    # 使用 SELECT ... FOR UPDATE 锁定 xiaoqi_new 记录
                    select_query = """
                                   SELECT xiaoqi_id
                                   FROM xiaoqi_new
                                   WHERE xiaoqi_name = %s
                                       FOR UPDATE \
                                   """
                    cursor.execute(select_query, (key,))
                    xiaoqi_record = cursor.fetchone()

                    if not xiaoqi_record:
                        raise ValueError(f"关联实体 {key} 不存在")

                    xiaoqi_id = xiaoqi_record[0]

                    # 批量插入 xiaoqi_to_file
                    insert_query = """
                                   INSERT INTO xiaoqi_to_file (xiaoqi_id, file_id)
                                   VALUES (%s, %s) ON DUPLICATE KEY \
                                   UPDATE xiaoqi_id = xiaoqi_id \
                                   """
                    # 使用 executemany 批量插入
                    cursor.executemany(
                        insert_query,
                        [(xiaoqi_id, fid) for fid in file_ids]
                    )
        db.connection.commit()
        logger.info(f"🔥 updata_to_mysql_new 批量提交成功")
        return entity_file

    except Exception as e:
        logger.error(f"数据库操作失败: {e}")
        return {}

    finally:
        db.close()  # 确保连接被关闭


def add_to_xiaoqi(file_text_list, file_id_list, entity, file_dict, file_dict_rev, userID, only_name):
    """
    处理实体到xiaoqi的添加操作

    Returns:
        dict: {"status": "success/error", "message": "详细信息", "data": 数据(可选)}
    """
    db = MySQLDatabase(
        host="114.213.234.179",
        user="koroot",
        password="DMiC-4092",
        database="db_hp"
    )
    result_list = []
    customize_content = []
    driver = GraphDatabase.driver("bolt://114.213.232.140:37687", auth=("neo4j", "123456"))
    info_list = []
    xiaoqi_id = 0
    customize_content_list = []
    # 手动管理Redis分布式锁，根据查询结果决定释放时机
    lock = get_distributed_lock(entity.rstrip('0123456789'), timeout=5)
    entity_found = False
    try:
        # 获取分布式锁 - 阻塞模式，使用Redis原生BLPOP等待
        logger.info(f"⏳ 线程 {threading.current_thread().name} 正在等待获取entity分布式锁: {entity}")
        lock.acquire(blocking=True)  # 使用Redis原生阻塞机制
        logger.info(f"🔒 线程 {threading.current_thread().name} 成功获取entity分布式锁: {entity}")

        entity_list, entity_dict = db.query_entities_by_name(entity)
        if len(entity_list) != 0:
            entity_found = True
            # 找到实体时，立即尝试释放锁
            try:
                lock.release()
                logger.info(f"🔓 线程 {threading.current_thread().name} 立即释放entity分布式锁 (找到实体): {entity}")
            except Exception as release_error:
                # 释放锁失败，立即停止并返回错误
                error_msg = f"释放分布式锁失败 - entity: {entity}, 错误: {str(release_error)}"
                logger.error(f"❌ {error_msg}")
                return {
                    "status": "error",
                    "message": error_msg,
                    "error_type": "lock_release_error",
                    "entity": entity,
                    "file_id_list": file_id_list
                }

        try:
            db.connect()
            for file_text, file_id in zip(file_text_list, file_id_list):
                file_data = {
                    'entity': entity.rstrip('0123456789'),
                    'sim': 1,
                    'file_id': file_id,
                }
                db.insert_data_without_primary("entity_to_file", file_data)

            if entity_found:
                result = {}
                for file_text, file_id in zip(file_text_list, file_id_list):
                    dictRes = xiaoqi_instance(file_text, entity_list)
                    info = db.query_dir_by_name_id(entity_dict[dictRes["entity"]])
                    customize_content = db.get_dir_private_list(entity_dict[dictRes["entity"]], userID)

                    if dictRes["entity"] in result:
                        existing_keywords = set(result[dictRes["entity"]][0])
                        new_keywords = set(dictRes["entity_with_keyword"].split(','))
                        combined_keywords = list(existing_keywords.union(new_keywords))

                        result[dictRes["entity"]][0] = combined_keywords
                        result[dictRes["entity"]][1].append(file_id)
                    else:
                        result[dictRes["entity"]] = [dictRes["entity_with_keyword"].split(','), [file_id]]

                    info_list.append(info)
                    customize_content_list.append(customize_content)
                entity_dict = updata_to_mysql_new(result, True)
                logger.info(f"✅ 实体处理完成 (锁已释放): {entity}")
            else:
                result, _, _ = jiekou_3(entity, userID)
                entity_dict = updata_to_mysql_new(result, True)
                file_dict = db.search_file(result, file_dict, file_dict_rev)
                logger.info(f"⏳ 线程 {threading.current_thread().name} 保持entity分布式锁 (未找到实体): {entity}")

            final_output = generate_entity_json(result, file_text_list, file_dict, info_list, customize_content_list,
                                                only_name, file_id_list)

            try:
                from query_neo4j.rabbitmq_client_producer import send_classification_tasks_and_wait

                # 异步发送分类任务到RabbitMQ并等待完成
                success = send_classification_tasks_and_wait(
                    final_output=final_output,
                    file_dict_rev=file_dict_rev,
                    entity_id=entity_dict[list(entity_dict.keys())[0]] if entity_dict else None,
                    user_id=userID,
                    xiaoqi_name=entity
                )

                if success:
                    logger.info(f"✅ 所有分类任务已完成 - entity: {entity}, payload_count: {len(final_output)}")
                else:
                    logger.warning(f"⚠️ 分类任务部分失败或超时 - entity: {entity}, payload_count: {len(final_output)}")

                logger.info(f"✅ 文件上传、实体创建和分类任务处理完成 - entity: {entity}, file_id: {file_id}")

            except Exception as async_e:
                logger.error(f"❌ 分类任务处理失败，错误: {str(async_e)}")
                # 即使分类失败，实体和文件已经成功创建，返回部分成功状态
                logger.info(f"ℹ️ 实体和文件创建成功，但分类任务失败 - entity: {entity}, file_id: {file_id}")

        except Exception as inner_e:
            error_msg = f"内部操作失败 - entity: {entity}, file_id: {file_id}, 错误: {str(inner_e)}"
            # 发生异常时尝试释放锁
            if not entity_found:
                try:
                    lock.release()
                    logger.info(f"🔓 线程 {threading.current_thread().name} 异常释放entity分布式锁: {entity}")
                except Exception as release_error:
                    # 释放锁失败，立即停止并返回错误
                    lock_error_msg = f"异常处理中释放分布式锁失败 - entity: {entity}, 原始错误: {error_msg}, 释放锁错误: {str(release_error)}"
                    logger.error(f"❌ {lock_error_msg}")
                    return {
                        "status": "error",
                        "message": lock_error_msg,
                        "error_type": "lock_release_error_in_exception",
                        "entity": entity,
                        "file_id_list": file_id_list,
                        "original_error": error_msg
                    }

            return {
                "status": "error",
                "message": error_msg,
                "error_type": "inner_operation_error",
                "entity": entity,
                "file_id_list": file_id_list
            }

    except Exception as outer_e:
        error_msg = str(outer_e)
        detailed_msg = f"entity分布式锁操作异常: {error_msg}"
        logger.error(f"❌ {detailed_msg}")

        # 如果是Redis连接问题，提供更具体的错误信息
        if "Redis" in error_msg or "Connection" in error_msg:
            detailed_msg += " (Redis服务器可能未启动或连接配置有问题)"
            logger.warning(f"💡 Redis服务器可能未启动或连接配置有问题")
            logger.debug(f"💡 Entity原文: {entity}")

        # 确保锁被释放
        try:
            lock.release()
            logger.info(f"🔓 线程 {threading.current_thread().name} 外层异常释放entity分布式锁: {entity}")
        except Exception as release_error:
            # 释放锁失败，立即停止并返回错误
            lock_error_msg = f"外层异常处理中释放分布式锁失败 - entity: {entity}, 原始错误: {detailed_msg}, 释放锁错误: {str(release_error)}"
            logger.error(f"❌ {lock_error_msg}")
            return {
                "status": "error",
                "message": lock_error_msg,
                "error_type": "lock_release_error_in_outer_exception",
                "entity": entity,
                "file_id_list": file_id_list,
                "original_error": detailed_msg
            }

        return {
            "status": "error",
            "message": detailed_msg,
            "error_type": "redis_lock_error",
            "entity": entity,
            "file_id_list": file_id_list
        }

    finally:
        # 如果没找到entity，在这里最终释放锁
        if not entity_found:
            try:
                lock.release()
                logger.info(f"🔓 线程 {threading.current_thread().name} 最终释放entity分布式锁 (函数结束): {entity}")
            except Exception as release_error:
                # 在finally块中无法通过return停止，但可以抛出异常
                final_error_msg = f"最终释放分布式锁失败 - entity: {entity}, 错误: {str(release_error)}"
                logger.error(f"❌ {final_error_msg}")
                # 抛出异常，这会被外层的process_file捕获
                raise Exception(final_error_msg)
        else:
            logger.debug(f"🏁 函数结束 - entity: {entity} (锁已在前面释放)")

    # 成功执行到这里，返回成功状态
    return {
        "status": "success",
        "message": f"实体处理成功 - entity: {entity}, file_id_list: {file_id_list}",
        "entity": entity,
        "file_id_list": file_id_list,
        "data": {
            "result": result,
            "entity_found": entity_found,
            "final_output": locals().get('final_output', [])
        }
    }


def create_entity_and_link(tx, name, path, userid):
    import datetime
    current_date = datetime.datetime.now()
    year = current_date.year
    month = current_date.month
    day = current_date.day
    minute = current_date.minute
    formatted_date = f"{year}-{month}-{day}"
    formatted_min = f"{year}-{month}-{day}_{minute}"

    query = """
            MERGE (f:Strict {name: $filename, path: $path, private: $private, user_id: $user_id, timestamp: $time})
            WITH f
            RETURN id(f) as id
            """
    # RETURN id(f) AS file_id, id(h) AS hypernode_id, id(k) AS category_id
    result = tx.run(
        query,
        filename=name,
        path=path,
        private=1,
        user_id=userid,
        time=formatted_date,
    )
    record = result.single()
    return record["id"]


def create_strict_and_coaser(tx, id, direct):
    query = """
        MATCH (a:Strict) WHERE id(a) = $id    
        MATCH (b:coarse {name: $direct})       
        MERGE (a)-[r:edge]->(b)
         RETURN r IS NOT NULL AS exists   
        """
    # RETURN id(f) AS file_id, id(h) AS hypernode_id, id(k) AS category_id
    result = tx.run(
        query,
        id=id,
        direct=direct
    )
    return True


def search_File_name(tx, name):
    query_word = ("MATCH (h:Strict) WHERE h.name = $name RETURN h")
    result = tx.run(query_word, name=name)
    return result.data()


def add_to_mysql(result_list, file_id_list, private, url_list=None):
    db = MySQLDatabase(
        host="114.213.234.179",
        user="koroot",  # 替换为您的用户名
        password="DMiC-4092",  # 替换为您的密码
        database="db_hp"  # 替换为您的数据库名
    )
    for result, file_id, url in zip(result_list, file_id_list, url_list):
        try:
            db.connect()
            data_to_insert = {
                "id": str(file_id),  # 假设 id 是自增字段，可以设置为 None
                "name": result[0]['h']["name"],  # 替换为实际数据
                "path": result[0]['h']["path"],  # 替换为实际时间戳，格式：YYYY-MM-DD HH:MM:SS
                "timestamp": result[0]['h']["timestamp"],  # 替换为实际 URL
                "private": private,
                "userid": result[0]['h']["user_id"],
                "url": url if url else None  # 添加URL字段，如果没有URL则为空字符串
            }

            db.insert_data("file", data_to_insert)
        finally:
            # 关闭数据库连接
            db.close()


def process_single_file_thread_safe(file_name_original, file_path, userid, index):
    """
    线程安全的单文件处理函数：Minio上传 + 文件解析 + Neo4j操作

    Args:
        file_name_original: 原始文件名
        file_path: 文件路径
        userid: 用户ID
        index: 文件索引

    Returns:
        Dict包含处理结果或错误信息
    """
    thread_name = threading.current_thread().name
    try:
        logger.info(f"🚀 [{thread_name}] 开始处理文件 [{index}]: {file_name_original}")

        # Minio上传步骤
        minio_address = "114.213.232.140:19000"
        minio_admin = "minioadmin"
        minio_password = "minioadmin"
        bucket = Bucket(minio_address=minio_address,
                        minio_admin=minio_admin,
                        minio_password=minio_password)

        # 生成文件信息
        doc_type = file_name_original.split('.')[-1].lower()
        file_name = file_name_original.split('.')[0] + '_' + str(int(time.time() * 1000)) + '.' + \
                    doc_type
        path = get_sha1_hash('upload')[:2] + '/' + file_name

        # Minio上传
        logger.info(f"📤 [{thread_name}] 开始Minio上传 [{index}]: {file_name}")
        bucket.upload_file_to_bucket('kofiles', path, file_path)
        logger.info(f"✅ [{thread_name}] Minio上传成功 [{index}]: {file_name}")

        # 步骤1：文件内容解析
        logger.info(f"🔍 [{thread_name}] 开始文件解析 [{index}]: {file_path} (类型: {doc_type})")
        doc = extract_text_with_markitdown_safe(file_path, file_name_original, doc_type)

        # 检查解析后的内容
        if not doc or not doc.strip():
            logger.warning(f"⚠️ [{thread_name}] 文件解析后内容为空，使用文件名 [{index}]: {file_path}")
            doc = file_name_original
        elif doc == file_name_original:
            logger.info(f"📄 [{thread_name}] 使用文件名作为内容 [{index}]: {file_path}")
        else:
            logger.info(f"✅ [{thread_name}] 文件解析成功 [{index}]: {file_path}, 内容长度: {len(str(doc))} 字符")

        # 步骤2：Neo4j操作
        logger.info(f"🔗 [{thread_name}] 开始Neo4j操作 [{index}]: {file_name}")
        driver = GraphDatabase.driver("bolt://114.213.232.140:37687", auth=("neo4j", "123456"))
        with driver.session() as session:
            file_id = session.write_transaction(
                create_entity_and_link,
                file_name,
                path,
                userid
            )
            search_result = session.write_transaction(search_File_name, file_name)

        logger.info(f"✅ [{thread_name}] Neo4j操作成功 [{index}]: file_id={file_id}")

        # 返回成功结果
        return {
            "status": "success",
            "index": index,
            "data": {
                "file_name_original": file_name_original,
                "file_name": file_name,
                "file_path": file_path,
                "path": path,
                "doc_type": doc_type,
                "content": doc,
                "file_id": file_id,
                "search_result": search_result,
                "thread_name": thread_name
            }
        }

    except Exception as e:
        error_message = f"文件处理失败 [{index}]: {str(e)}"
        logger.error(f"❌ [{thread_name}] {error_message}")

        # 检查错误类型
        if "Minio" in str(e) or "upload" in str(e).lower():
            error_type = "minio_upload_error"
        elif "Neo4j" in str(e) or "ServiceUnavailable" in str(e):
            error_type = "neo4j_error"
        else:
            error_type = "general_error"

        return {
            "status": "error",
            "index": index,
            "message": error_message,
            "error_type": error_type,
            "file_info": {
                "file_name_original": file_name_original,
                "file_path": file_path,
                "thread_name": thread_name
            }
        }


def process_file(file_name_original_list, file_path_list, name, userid, private, url_list=None, only_name=False):
    """
    处理文件上传和实体提取
    Returns:
        dict: {"status": "success/error", "message": "详细信息", "data": 数据(可选)}
    """
    try:
        path_list = []
        doc_list = []
        doc_type_list = []
        file_name_list = []
        file_id_list = []
        file_dict = {}
        file_dict_rev = {}
        result_list = []
        # 确保url_list不为None，如果为None则创建相同长度的None列表
        if url_list is None:
            url_list = [None] * len(file_name_original_list)

        max_workers = min(4, len(file_name_original_list))  # 限制最大线程数
        logger.info(f"🔧 启动多线程处理，文件数量: {len(file_name_original_list)}, 线程数: {max_workers}")

        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            # 提交所有文件处理任务
            future_to_index = {}
            for i, (file_name_original, file_path) in enumerate(zip(file_name_original_list, file_path_list)):
                future = executor.submit(process_single_file_thread_safe, file_name_original, file_path, userid, i)
                future_to_index[future] = i

            # 初始化结果存储（按索引顺序）
            results = [None] * len(file_name_original_list)

            # 收集结果 - 等待所有任务完成
            failed_tasks = []
            for future in as_completed(future_to_index):
                index = future_to_index[future]
                try:
                    result = future.result()
                    results[index] = result

                    if result["status"] == "error":
                        logger.error(f"❌ 多线程任务失败 [{index}]: {result['message']}")
                        failed_tasks.append({"index": index, "error": result})
                    else:
                        logger.info(f"✅ 多线程任务完成 [{index}]: {result['data']['file_name_original']}")

                except Exception as task_e:
                    error_message = f"多线程任务异常 [{index}]: {str(task_e)}"
                    logger.error(f"❌ {error_message}")
                    error_result = {
                        "status": "error",
                        "message": error_message,
                        "error_type": "thread_execution_error",
                        "file_info": {"index": index}
                    }
                    results[index] = error_result
                    failed_tasks.append({"index": index, "error": error_result})

            # 如果有失败的任务，返回第一个错误
            if failed_tasks:
                first_error = failed_tasks[0]["error"]
                return {
                    "status": "error",
                    "message": first_error["message"],
                    "error_type": first_error.get("error_type", "thread_error"),
                    "file_info": first_error.get("file_info", {}),
                    "failed_count": len(failed_tasks),
                    "total_count": len(file_name_original_list)
                }

        # 检查是否所有任务都完成
        if None in results:
            error_message = "部分多线程任务未完成"
            logger.error(f"❌ {error_message}")
            return {
                "status": "error",
                "message": error_message,
                "error_type": "incomplete_threads_error"
            }

        # 按顺序整理结果，只处理成功的任务
        logger.info(f"📋 整理多线程处理结果，共 {len(results)} 个文件")
        successful_indices = []  # 记录成功任务的原始索引

        for i, result in enumerate(results):
            # 只处理成功的任务
            if result["status"] == "success":
                data = result["data"]

                doc_list.append(data["content"])
                file_id_list.append(data["file_id"])
                file_name_list.append(data["file_name"])
                path_list.append(data["path"])
                doc_type_list.append(data["doc_type"])

                file_dict[data["file_id"]] = f"bb/{data['file_name']}"
                file_dict_rev[f"bb/{data['file_name']}"] = data["file_id"]
                result_list.append(data["search_result"])

                successful_indices.append(i)  # 记录成功的原始索引
                logger.debug(
                    f"✅ 结果整理完成 [{i}]: file_id={data['file_id']}, content_length={len(str(data['content']))}")
            else:
                logger.warning(f"⚠️ 跳过失败任务 [{i}]: {result['message']}")

        if url_list:
            url_list = [url_list[i] if i < len(url_list) else None for i in successful_indices]
        else:
            url_list = [None] * len(successful_indices)

        logger.info(
            f"🎯 多线程处理和结果整理完成，成功处理 {len(successful_indices)}/{len(file_name_original_list)} 个文件")
        logger.info(f"📊 成功任务索引顺序: {successful_indices}")

        if not successful_indices:
            return {
                "status": "error",
                "message": "所有文件处理任务都失败了",
                "error_type": "all_tasks_failed",
                "total_count": len(file_name_original_list)
            }

        # 验证列表长度一致性
        if len(doc_list) != len(file_id_list):
            logger.error(f"❌ 严重错误：doc_list长度({len(doc_list)}) != file_id_list长度({len(file_id_list)})")
            return {
                "status": "error",
                "message": "文件内容和文件ID列表长度不匹配",
                "error_type": "data_consistency_error"
            }

        logger.info(f"✅ 文件处理完成，成功处理 {len(file_id_list)} 个文件，doc_list和file_id_list顺序一致")
        # 调用 add_to_mysql 并处理可能的异常
        try:
            add_to_mysql(result_list, file_id_list, private, url_list)
        except Exception as mysql_e:
            error_message = f"MySQL数据库操作失败: {str(mysql_e)}"
            logger.error(f"❌ MySQL操作异常: {error_message}")
            return {
                "status": "error",
                "message": error_message,
                "error_type": "mysql_operation_error",
                "file_info": {
                    "file_name_original": file_name_original,
                    "file_id_list": file_id_list,
                    "path_list": path_list
                }
            }

        # 调用 add_to_xiaoqi 并处理返回结果
        # 传入file_id_list确保正确的文件内容映射
        xiaoqi_result = add_to_xiaoqi(doc_list, file_id_list, name, file_dict, file_dict_rev, userid, only_name)

        # 检查 add_to_xiaoqi 的执行结果
        if xiaoqi_result["status"] == "error":
            # 如果 add_to_xiaoqi 执行失败，返回错误信息
            return {
                "status": "error",
                "message": f"实体处理失败: {xiaoqi_result['message']}",
                "error_details": xiaoqi_result,
                "file_info": {
                    "file_name_original": file_name_original,
                    "file_id_list": file_id_list,
                    "entity": name
                }
            }

        return {
            "status": "success",
            "message": "文件处理完成",
            "file_info": {
                "file_name_original": file_name_original,
                "file_id_list": file_id_list,
                "entity": name,
                "path_list": path_list
            },
            "xiaoqi_result": xiaoqi_result
        }

    except Exception as process_e:
        # 捕获所有其他异常
        import traceback
        error_msg = f"process_file执行失败: {str(process_e)}"
        error_trace = traceback.format_exc()
        logger.error(f"❌ {error_msg}")
        logger.debug(f"异常堆栈: {error_trace}")

        return {
            "status": "error",
            "message": error_msg,
            "error_type": "process_file_error",
            "error_trace": error_trace,
            "file_info": {
                "file_name": file_name_original,
                "file_path": file_path,
                "entity": name,
                "userid": userid
            }
        }


def main(request):
    try:
        # 从GET请求中获取参数
        name = request.GET.get("name")
        only_name = request.GET.get("only_name", "false").lower() == "true"
        userid = request.GET.get("userid")
        # remote_path = request.GET.get('path', None)
        remote_path = request.GET.get('path_list', None)
        private = int(request.GET.get('private', 1))
        url_list = request.GET.get('url_list', None)
    except Exception as param_e:
        return JsonResponse({"status": "error", "message": f"参数解析失败: {str(param_e)}"}, status=400)

    # 验证必要参数
    if not name:
        return JsonResponse({"status": "error", "message": "缺少必要参数'name'"}, status=400)
    if not userid:
        return JsonResponse({"status": "error", "message": "缺少必要参数'userid'"}, status=400)

    try:
        userid = int(userid)
        private = int(private)
    except ValueError:
        return JsonResponse({"status": "error", "message": "userid和private必须是整数"}, status=400)

    head_path = 'D:/upload/'
    file_path_list = []
    file_name_original_list = []
    try:
        if remote_path:
            # 使用远程服务器提供的路径
            for path in remote_path:
                file_name_original = os.path.basename(path)
                file_path = os.path.join(head_path, file_name_original)
                file_name_original_list.append(file_name_original)
                file_path_list.append(file_path)
            # 从远程路径直接读取文件，不需要从file_obj.chunks()读取
        else:
            # 原有逻辑：从上传的文件对象读取
            file_obj = request.FILES.get('file', None)

            if not file_obj:
                return JsonResponse({
                    "status": "error",
                    "message": "未提供文件",
                    "error_type": "file_missing_error"
                }, status=400)

            file_path = os.path.join(head_path, file_obj.name)
            with open(file_path, 'wb') as f:
                for chunk in file_obj.chunks():
                    f.write(chunk)
            file_name_original_list.append(file_obj.name)
            file_path_list.append(file_path)

    except IOError as io_e:
        return JsonResponse({
            "status": "error",
            "message": f"文件操作失败: {str(io_e)}",
            "error_type": "file_operation_error"
        }, status=500)
    except Exception as file_e:
        return JsonResponse({
            "status": "error",
            "message": f"文件处理失败: {str(file_e)}",
            "error_type": "file_handling_error"
        }, status=500)
    try:
        result = process_file(file_name_original_list, file_path_list, name, userid, private, url_list, only_name)

        # 检查 process_file 的返回结果
        if isinstance(result, dict) and result.get("status") == "error":
            # process_file 返回了错误信息
            error_type = result.get("error_type", "unknown_error")

            # 根据错误类型设置不同的HTTP状态码
            if error_type in ["neo4j_connection_error", "neo4j_operation_error"]:
                status_code = 503  # Service Unavailable - 数据库服务不可用
            elif error_type in ["redis_lock_error", "lock_release_error", "lock_release_error_in_exception",
                                "lock_release_error_in_outer_exception"]:
                status_code = 503  # Service Unavailable - Redis锁服务不可用
            elif error_type == "mysql_operation_error":
                status_code = 503  # Service Unavailable - MySQL服务不可用
            elif error_type == "minio_upload_error":
                status_code = 502  # Bad Gateway - 文件存储服务异常
            elif error_type == "process_file_error":
                status_code = 500  # Internal Server Error
            else:
                status_code = 400  # Bad Request

            return JsonResponse({
                "status": "error",
                "message": result["message"],
                "error_type": error_type,
                "error_details": result.get("error_details", {}),
                "file_info": result.get("file_info", {}),
                "error_trace": result.get("error_trace", "")
            }, status=status_code)
        elif isinstance(result, str) and "error" in result:
            # 处理旧的JSON字符串格式错误
            return JsonResponse({"status": "error", "message": "上传文件失败", "details": result}, status=500)
        else:
            # 成功处理
            return JsonResponse({
                "status": "success",
                "message": "文件处理成功",
                "data": result
            }, safe=False)

    except Exception as e:
        import traceback
        error_msg = str(e)
        error_trace = traceback.format_exc()
        logger.error(f"main函数异常: {error_msg}")
        logger.debug(f"异常堆栈: {error_trace}")
        return JsonResponse({
            "status": "error",
            "message": f"处理文件时出错: {error_msg}",
            "error_trace": error_trace
        }, status=500)


def extract_text_with_pymupdf(pdf_path):
    """使用PyMuPDF解析PDF，带有详细错误处理和备选方案"""
    try:
        # 首先尝试使用PyMuPDF解析
        logger.info(f": {pdf_path}")
        doc = fitz.open(pdf_path)
        text = ""

        # 逐页提取文本，并处理单页错误
        total_pages = 1
        successful_pages = 0

        for page_num in range(total_pages):
            try:
                page = doc.load_page(page_num)
                page_text = page.get_text()
                text += page_text
                successful_pages += 1
            except Exception as page_e:
                logger.warning(f"⚠️ PDF第{page_num + 1}页解析失败: {str(page_e)}")
                # 继续处理下一页
                return None

        doc.close()

        logger.info(f"✅ PDF解析完成: 总页数={total_pages}, 成功页数={successful_pages}")

        # 检查是否提取到有效内容
        if text and text.strip():
            logger.debug(f"📝 提取文本长度: {len(text)} 字符")
            return text
        else:
            logger.warning("⚠️ 未提取到文本内容，尝试备选方案")
            return None

    except Exception as mupdf_e:
        error_msg = str(mupdf_e)
        logger.error(f"❌ PyMuPDF解析失败: {error_msg}")

        # 检查是否是常见的MuPDF错误
        if any(keyword in error_msg.lower() for keyword in [
            'syntax error', 'content stream', 'invalid key',
            'expected object', 'damaged', 'corrupt'
        ]):
            logger.info("🔍 检测到PDF格式问题，尝试备选解析方案...")

            # 尝试备选方案1：使用pdfminer
            try:
                logger.info("🔄 尝试使用pdfminer解析...")
                from pdfminer.high_level import extract_text
                text = extract_text(pdf_path)
                if text and text.strip():
                    logger.info(f"✅ pdfminer解析成功，提取文本长度: {len(text)} 字符")
                    return text
            except Exception as pdfminer_e:
                logger.warning(f"❌ pdfminer解析也失败: {str(pdfminer_e)}")

            # 尝试备选方案2：使用PyPDF2
            try:
                logger.info("🔄 尝试使用PyPDF2解析...")
                import PyPDF2
                with open(pdf_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    text = ""
                    for page_num in range(len(pdf_reader.pages)):
                        try:
                            page = pdf_reader.pages[page_num]
                            text += page.extract_text()
                        except Exception as page_e:
                            logger.warning(f"⚠️ PyPDF2第{page_num + 1}页解析失败: {str(page_e)}")
                            continue

                    if text and text.strip():
                        logger.info(f"✅ PyPDF2解析成功，提取文本长度: {len(text)} 字符")
                        return text
            except Exception as pypdf2_e:
                logger.warning(f"❌ PyPDF2解析也失败: {str(pypdf2_e)}")

            logger.warning("❌ 所有PDF解析方案都失败，返回None")
            return None
        else:
            # 其他类型的错误，直接返回None
            logger.error(f"❌ 其他类型的PDF解析错误: {error_msg}")
            return None


class GenericHtmlParser(HTMLParser):
    """
    通用HTML解析器，提取所有可见文本内容
    """

    def __init__(self):
        super().__init__()
        self.text_content = []
        self.current_text = ""
        self.skip_tags = {'script', 'style', 'head', 'title', 'meta', 'link'}
        self.in_skip_tag = False

    def handle_starttag(self, tag, attrs):
        if tag.lower() in self.skip_tags:
            self.in_skip_tag = True

    def handle_endtag(self, tag):
        if tag.lower() in self.skip_tags:
            self.in_skip_tag = False
        elif not self.in_skip_tag and self.current_text.strip():
            # 在块级元素结束时保存文本
            if tag.lower() in {'p', 'div', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'li', 'td', 'th'}:
                self.text_content.append(self.current_text.strip())
                self.current_text = ""

    def handle_data(self, data):
        if not self.in_skip_tag:
            self.current_text += data

    def get_text_content(self):
        # 添加最后的文本内容
        if self.current_text.strip():
            self.text_content.append(self.current_text.strip())

        # 过滤空内容和过短内容
        filtered_content = [content for content in self.text_content
                            if content and len(content.strip()) > 10]

        return filtered_content[:10] if filtered_content else ["HTML内容解析为空"]