import sys
import time
import hashlib
from django.http import HttpResponse, Http404, StreamingHttpResponse
from django.utils.encoding import escape_uri_path
import zipfile
import os
import tqdm
from minio import Minio, InvalidResponseError, S3Error
from neo4j import GraphDatabase
import zhconv
import shutil
import json
import chardet
import pymysql
import ctypes
from query_neo4j.search_urls import search_urls
import os
import re
import time
import json
import requests
from typing import Optional, Dict
from django.http import HttpRequest  # 添加这行导入
from django.http import HttpResponse as DjangoHttpResponse  # 如果需要

class MySQLDatabase:
    def __init__(self, host, user, password, database, charset="utf8mb4"):
        """
        初始化数据库连接
        """
        self.config = {
            "host": host,
            "user": user,
            "password": password,
            "database": database,
            "charset": charset
        }
        self.connection = None

    def connect(self):
        """
        建立数据库连接
        """
        try:
            self.connection = pymysql.connect(**self.config)
            print("数据库连接成功！")
        except pymysql.MySQLError as e:
            print(f"数据库连接失败：{e}")
            raise

    def insert_data(self, table_name, data):
        try:
            # 先检查主键是否存在
            primary_key = list(data.keys())[0]  # 假设主键在第一个位置
            primary_key_value = data[primary_key]

            # 生成检查主键是否存在的 SQL 查询
            check_query = f"SELECT COUNT(*) FROM {table_name} WHERE {primary_key} = %s"
            with self.connection.cursor() as cursor:
                cursor.execute(check_query, (primary_key_value,))
                result = cursor.fetchone()

                if result[0] > 0:
                    print(f"主键 {primary_key_value} 已存在，跳过插入操作。")
                    return  # 主键已存在，跳过插入操作

            # 生成插入 SQL 语句
            columns = ", ".join(data.keys())
            placeholders = ", ".join(["%s"] * len(data))
            insert_query = f"INSERT INTO {table_name} ({columns}) VALUES ({placeholders})"

            # 执行插入操作
            with self.connection.cursor() as cursor:
                cursor.execute(insert_query, tuple(data.values()))
                self.connection.commit()
                print("数据插入成功！")
        except pymysql.MySQLError as e:
            print(f"插入数据失败：{e}")
            self.connection.rollback()  # 回滚事务

    def insert_relation(self, table_name, data):
        try:
            # 先检查主键是否存在
            primary_key = list(data.keys())[0]  # 假设主键在第一个位置
            last_key = list(data.keys())[-1]
            primary_key_value = data[primary_key]
            last_key_value = data[last_key]

            # 生成检查主键是否存在的 SQL 查询
            check_query = f"SELECT COUNT(*) FROM {table_name} WHERE {primary_key} = %s AND {last_key} = %s "
            with self.connection.cursor() as cursor:
                cursor.execute(check_query, (primary_key_value, last_key_value,))
                result = cursor.fetchone()

                if result[0] > 0:
                    print(f"主键 {primary_key_value} 已存在，跳过插入操作。")
                    return  # 主键已存在，跳过插入操作

            # 生成插入 SQL 语句
            columns = ", ".join(data.keys())
            placeholders = ", ".join(["%s"] * len(data))
            insert_query = f"INSERT INTO {table_name} ({columns}) VALUES ({placeholders})"

            # 执行插入操作
            with self.connection.cursor() as cursor:
                cursor.execute(insert_query, tuple(data.values()))
                self.connection.commit()
                print("数据插入成功！")
        except pymysql.MySQLError as e:
            print(f"插入数据失败：{e}")
            self.connection.rollback()  # 回滚事务

    def query_tables(self, query):
        try:
            with self.connection.cursor() as cursor:
                cursor.execute(query)
                result = cursor.fetchall()
                return result
        except pymysql.MySQLError as e:
            print(f"查询失败：{e}")
            raise

    def close(self):
        """
        关闭数据库连接
        """
        if self.connection:
            self.connection.close()
            print("数据库连接已关闭！")


class Bucket:

    def __init__(self, minio_address, minio_admin, minio_password):
        # 通过ip 账号 密码 连接minio server
        # Http连接 将secure设置为False
        self.minioClient = Minio(endpoint=minio_address,
                                 access_key=minio_admin,
                                 secret_key=minio_password,
                                 secure=False)

    def create_one_bucket(self, bucket_name):
        # 创建桶(调用make_bucket api来创建一个桶)
        """
        桶命名规则：小写字母，句点，连字符和数字 允许使用 长度至少3个字符
        使用大写字母、下划线等会报错
        """
        try:
            # bucket_exists：检查桶是否存在
            if self.minioClient.bucket_exists(bucket_name=bucket_name):
                print("该存储桶已经存在")
            else:
                self.minioClient.make_bucket(bucket_name=bucket_name)
                print(f"{bucket_name}桶创建成功")
        except InvalidResponseError as err:
            print(err)

    def remove_one_bucket(self, bucket_name):
        # 删除桶(调用remove_bucket api来创建一个存储桶)
        try:
            if self.minioClient.bucket_exists(bucket_name=bucket_name):
                self.minioClient.remove_bucket(bucket_name)
                print("删除存储桶成功")
            else:
                print("该存储桶不存在")
        except InvalidResponseError as err:
            print(err)

    def upload_file_to_bucket(self, bucket_name, file_name, file_path):
        """
        将文件上传到bucket
        :param bucket_name: minio桶名称
        :param file_name: 存放到minio桶中的文件名字(相当于对文件进行了重命名，可以与原文件名不同)
                            file_name处可以创建新的目录(文件夹) 例如 /example/file_name
                            相当于在该桶中新建了一个example文件夹 并把文件放在其中
        :param file_path: 本地文件的路径
        """
        # 桶是否存在 不存在则新建
        check_bucket = self.minioClient.bucket_exists(bucket_name)
        if not check_bucket:
            self.minioClient.make_bucket(bucket_name)

        try:
            self.minioClient.fput_object(bucket_name=bucket_name,
                                         object_name=file_name,
                                         file_path=file_path)
        except FileNotFoundError as err:
            print('upload_failed: ' + str(err))
        except S3Error as err:
            print("upload_failed:", err)

    def download_file_from_bucket(self, bucket_name, minio_file_path, download_file_path):
        """
        从bucket下载文件
        :param bucket_name: minio桶名称
        :param minio_file_path: 存放在minio桶中文件名字
                            file_name处可以包含目录(文件夹) 例如 /example/file_name
        :param download_file_path: 文件获取后存放的路径
        """
        # 桶是否存在
        check_bucket = self.minioClient.bucket_exists(bucket_name)
        if check_bucket:
            try:
                self.minioClient.fget_object(bucket_name=bucket_name,
                                             object_name=minio_file_path,
                                             file_path=download_file_path)
                return 1
            except FileNotFoundError as err:
                print('download_failed: ' + str(err))
                return 0
            except S3Error as err:
                print("download_failed:", err)
                return 0

    def remove_object(self, bucket_name, object_name):
        """
        从bucket删除文件
        :param bucket_name: minio桶名称
        :param object_name: 存放在minio桶中的文件名字
                            object_name处可以包含目录(文件夹) 例如 /example/file_name
        """
        # 桶是否存在
        check_bucket = self.minioClient.bucket_exists(bucket_name)
        if check_bucket:
            try:
                self.minioClient.remove_object(bucket_name=bucket_name,
                                               object_name=object_name)
            except FileNotFoundError as err:
                print('upload_failed: ' + str(err))
            except S3Error as err:
                print("upload_failed:", err)

    # 获取所有的桶
    def get_all_bucket(self):
        buckets = self.minioClient.list_buckets()
        ret = []
        for _ in buckets:
            ret.append(_.name)
        return ret

    # 获取一个桶中的所有一级目录和文件
    def get_list_objects_from_bucket(self, bucket_name):
        # 桶是否存在
        check_bucket = self.minioClient.bucket_exists(bucket_name)
        if check_bucket:
            # 获取到该桶中的所有目录和文件
            objects = self.minioClient.list_objects(bucket_name=bucket_name)
            ret = []
            for _ in objects:
                ret.append(_.object_name)
            return ret

    # 获取桶里某个目录下的所有目录和文件
    def get_list_objects_from_bucket_dir(self, bucket_name, dir_name):
        # 桶是否存在
        check_bucket = self.minioClient.bucket_exists(bucket_name)
        if check_bucket:
            # 获取到bucket_所name桶中的dir_name下的有目录和文件
            # prefix 获取的文件路径需包含该前缀
            objects = self.minioClient.list_objects(bucket_name=bucket_name,
                                                    prefix=dir_name,
                                                    recursive=True)
            ret = []
            for obj in objects:
                object_name = obj.object_name
                # 获取对象的内容
                content = self.minioClient.get_object(bucket_name=bucket_name,
                                                      object_name=object_name)
                ret.append(content.data.decode())
            return ret


def get_sha1_hash(file_name):
    shal_hash = hashlib.sha1(file_name.encode()).hexdigest()
    return shal_hash


def file_iterator(file_path, chunk_size=512):
    """
    文件生成器,防止文件过大，导致内存溢出
    :param file_path: 文件绝对路径
    :param chunk_size: 块大小
    :return: 生成器
    """
    with open(file_path, mode='rb') as f:
        while True:
            c = f.read(chunk_size)
            if c:
                yield c
            else:
                break


def deletezip(directory):
    for filename in os.listdir(directory):
        # 判断是否以 .zip 结尾
        if filename.endswith('.zip'):
            # 构建完整的文件路径
            file_path = os.path.join(directory, filename)
            # 删除文件
            os.remove(file_path)


def zipDir(dirpath, outFullName):
    """
    压缩指定文件夹
    :param dirpath: 目标文件夹路径
    :param outFullName: 压缩文件保存路径+xxxx.zip
    :return: 无
    """
    zip = zipfile.ZipFile(outFullName, "a", zipfile.ZIP_DEFLATED)
    flag=False
    for path, dirnames, filenames in os.walk(dirpath):
        # 去掉目标跟路径，只对目标文件夹下边的文件及文件夹进行压缩
        fpath = dirpath
        # print(dirpath.split("\\")[-1])
        # zip.write(path, os.path.relpath(path, os.path.dirname(dirpath)))
        for filename in filenames:
            flag=True
            print(str(path)+" "+str(dirnames)+" "+str(filenames))
            zip.write(os.path.join(path, filename), os.path.join(fpath, filename))
    if flag==False:
        print(dirpath)
        zip.write(dirpath, dirpath)
        # zip.writestr(os.path.basename(dirpath) + '/.keep', '')
    zip.close()

def search_content_to_file(content, db):
    """
    根据给定的 content 查询 entity_to_file 表中的 file_id，
    并返回 file 表中所有匹配的行的信息。
    """
    query = """
    SELECT ef.sim, f.*
    FROM file AS f
    JOIN entity_to_file AS ef ON f.id = ef.file_id
    WHERE ef.entity = %s;
    """
    try:
        with db.connection.cursor() as cursor:
            cursor.execute(query, (content,))
            result = cursor.fetchall()
            return result
    except pymysql.MySQLError as e:
        print(f"查询失败：{e}")
        raise


def check_node_exists(name, node_id, userID, session):
    result = session.run(
        """
        MATCH (n)
        WHERE id(n) = $node_id and n.name=$name and n.user_id=$userID
        RETURN CASE WHEN COUNT(n) > 0 THEN 1 ELSE 0 END AS exists
        """,
        node_id=int(node_id),
        name=name,
        userID=int(userID)
    )
    return result.single()[0]


def check_node_strict(name, node_id, userID, session):
    result = session.run(
        """
        MATCH (n)
        WHERE id(n) = $node_id and n.name=$name and n.user_id=$userID
        RETURN CASE WHEN COUNT(n) > 0 THEN 1 ELSE 0 END AS exists
        """,
        node_id=int(node_id),
        name=name,
        userID=int(userID)
    )
    return result.single()[0]
    # return result.single()[0]  # 返回查询结果的第一个值


def chuli_file_dire(under_dire, under_dire_to_file, session):
    for i in under_dire:
        query_word = '''
            MATCH (h:KOCategory)<-[r]-(t)
    WHERE id(h) = {arg_1} AND t:KOCategory
    RETURN h AS h, r AS r, t AS t, NULL AS fullPath, 0 AS pathLength
    LIMIT 100

    UNION ALL

    MATCH (h:KOCategory)
    WHERE id(h) = {arg_1}
    CALL apoc.path.expandConfig(h, {
        relationshipFilter: "<edge",
        labelFilter: ">KOCategory",
        maxLevel: 5,
        limit: 300,
        direction: "incoming"
    })
    YIELD path
    WITH h, path, relationships(path) AS rels  // 保留 h 和 path
    UNWIND rels AS rel
    WITH h, rel, startNode(rel) AS t, endNode(rel) AS ko, path, length(path) AS pathLength  // 计算路径长度
    WHERE ko:KOCategory AND t <> h  // 确保 t 不是 h
    RETURN h AS h, rel AS r, ko AS t, path AS fullPath, pathLength
    ORDER BY pathLength DESC  // 按路径长度降序排列
    LIMIT 1  // 仅返回最长的路径
        '''
        results_place = session.run(query_word, parameters={"arg_1": int(i)})
        dire_relation = []
        for res in results_place:
            if res["fullPath"] is not None:
                for edge in res["fullPath"]:
                    dire_relation.append((edge.start_node.id, edge.end_node.id))
                    # print(node)
        for pair in dire_relation:
            if int(pair[0]) not in under_dire_to_file:
                temp = []
            else:
                temp = under_dire_to_file[int(pair[0])]
            temp.append(int(pair[1]))
            under_dire_to_file[int(pair[0])] = temp
    return under_dire_to_file


def check_node_name(session, i):
    node1 = id_get_node(session, i)
    if "wikipage" in node1.labels:
        node_name = node1["name"] + "_维基.html"
    elif "baidupage" in node1.labels:
        node_name = node1["name"] + "_百度.html"
    else:
        node_name = node1["name"]
    return node_name


def id_get_node(session, id):
    result = session.run(
        """
        MATCH (n)
        WHERE id(n) = $node_id
        RETURN n
        """,
        node_id=int(id)
    )
    # print(id)
    return result.single()["n"]

def compress_file(file_path, zip_path):
    # 检查指定的文件是否存在
    if not os.path.isfile(file_path):
        print(f"文件 {file_path} 不存在。")
        return

    # 创建一个压缩包并将文件添加到其中
    with zipfile.ZipFile(zip_path, 'a') as zipf:
        zipf.write(file_path, os.path.basename(file_path))
        print(f"文件 {file_path} 已成功压缩到 {zip_path}。")

def search_fileId_in_mysql(db,fileid,userID):
    query = """
SELECT
    de.dir_private,
    f.name
FROM
    dir_file AS df
JOIN
    dir_entity AS de ON df.dir_id = de.id and de.userID=%s
JOIN
    file AS f ON df.file_id = f.id
WHERE
    df.file_id = %s
        """
    try:
        with db.connection.cursor() as cursor:
            cursor.execute(query, (int(userID),int(fileid),))
            result = cursor.fetchall()
            return result
    except pymysql.MySQLError as e:
        print(f"查询失败：{e}")
        raise
def search_dire_in_mysql(db,content,userid):
    query = """
    SELECT
    de.dir_private
FROM
    xiaoqi_new AS xn
JOIN
    dir_entity AS de ON xn.xiaoqi_id = de.entity_id and de.userid=%s
WHERE
    xn.xiaoqi_name=%s
            """
    try:
        with db.connection.cursor() as cursor:
            cursor.execute(query, (int(userid),str(content)))
            result = cursor.fetchall()
            return result
    except pymysql.MySQLError as e:
        print(f"查询失败：{e}")
        raise
def search_nodeIds(userID,content):
    db = MySQLDatabase(
            host="114.213.234.179",
            user="koroot",  # 替换为您的用户名
            password="DMiC-4092",  # 替换为您的密码
            database="db_hp"  # 替换为您的数据库名
        )
    db.connect()
    query = """
    SELECT
            xn.xiaoqi_id,
        de.id AS directory_id,
        de.dir_private,
        x.file_id AS xiaoqi_file_id,
            f.name AS file_name,
            f.private As file_private
    FROM
        xiaoqi_new xn
    JOIN
        xiaoqi_to_file x ON xn.xiaoqi_id = x.xiaoqi_id
    JOIN
        dir_file df ON x.file_id = df.file_id
    JOIN
        dir_entity de ON df.dir_id = de.id and de.entity_id=xn.xiaoqi_id and de.userid=%s
    JOIN
        file f ON x.file_id = f.id and (f.private=0 or f.userid=%s)
    WHERE
        xn.xiaoqi_name = %s;
        """
    try:
        with db.connection.cursor() as cursor:
            cursor.execute(query, (int(userID), int(userID), str(content)))
            result = cursor.fetchall()
    except pymysql.MySQLError as e:
        print(f"查询失败：{e}")
        raise
    nodes=[]
    for i in result:
        nodes.append(int(i[3]))
    return nodes

def test(content, userID):
    print(f"[DEBUG] 开始处理: {content}, 用户ID: {userID}")

    base_dir = os.path.join('D:\\data\\Auto_recommendtion', content)
    print(f"[DEBUG] 基础目录: {base_dir}")

    if not os.path.exists(base_dir):
        os.makedirs(base_dir)
        print(f"[DEBUG] 创建目录: {base_dir}")
    else:
        print(f"[DEBUG] 目录已存在: {base_dir}")

    # 数据库连接
    print("[DEBUG] 连接数据库...")
    db = MySQLDatabase(host="114.213.234.179", user="koroot", password="DMiC-4092", database="db_hp")
    db.connect()

    # 获取节点ID
    print("[DEBUG] 查询节点ID...")
    node_id = search_nodeIds(userID, content)
    print(f"[DEBUG] 找到节点ID: {node_id}")

    # 初始化 MinIO bucket - 确保这个在函数内部定义
    print("[DEBUG] 初始化MinIO连接...")
    bucket = Bucket(
        minio_address="114.213.232.140:19000",
        minio_admin="minioadmin",
        minio_password="minioadmin"
    )

    # 检查 bucket 连接
    try:
        buckets = bucket.get_all_bucket()
        print(f"[DEBUG] 可用的buckets: {buckets}")
        if 'kofiles' not in buckets:
            print("[ERROR] kofiles bucket 不存在!")
            return []
    except Exception as e:
        print(f"[ERROR] MinIO连接失败: {e}")
        return []

    # 创建目录
    result_dire = search_dire_in_mysql(db, content, userID)
    print(f"[DEBUG] 数据库目录结果: {result_dire}")

    for i in result_dire:
        dir_path = os.path.join(base_dir, i[0])
        print(f"[DEBUG] 处理目录: {dir_path}")
        if not os.path.exists(dir_path):
            os.makedirs(dir_path)
            print(f"[DEBUG] 创建子目录: {dir_path}")

    # 下载文件
    address = []
    for i in node_id:
        print(f"[DEBUG] 处理节点 {i}...")
        result = search_fileId_in_mysql(db, i, userID)
        print(f"[DEBUG] 文件查询结果: {result}")

        for res in result:
            dire_name = res[0]
            node_path = "bb/" + res[1]
            target_dir = os.path.join(base_dir, dire_name)
            filename = node_path.split('/')[-1]
            full_path = os.path.join(target_dir, filename)

            print(f"[DEBUG] 下载: {node_path} -> {full_path}")

            # 确保目标目录存在
            if not os.path.exists(target_dir):
                os.makedirs(target_dir)
                print(f"[DEBUG] 创建下载目录: {target_dir}")

            # 使用上面定义的 bucket 变量
            k = bucket.download_file_from_bucket('kofiles', node_path, full_path)
            if k == 1:
                print(f"[DEBUG] 下载成功: {full_path}")
                address.append(full_path)
            else:
                print(f"[ERROR] 下载失败: {node_path}")

    print(f"[DEBUG] 最终文件列表: {address}")
    return address

##########################################################################################################################################################################################
# print(test("吴信东1",6000622))#第一个参数代表消歧实体名称，对应xiaoqi_new表中的name，第二个是用户ID（这个到时候前端也会提供）
#第一个模块：前端传入 实体名 和 对应ID,调用test返回下载在本地文件的地址（test 函数通过 数据库查询 → 本地目录创建 → MinIO 文件下载 的流程）
#######################################
#######################################

#######################################
import os
import re
import pdfplumber
from docx import Document
from bs4 import BeautifulSoup
import requests
from typing import List, Optional
from concurrent.futures import ThreadPoolExecutor, as_completed
import chardet


def read_file_content(filepath: str) -> str:
    try:
        if filepath.endswith(".pdf"):
            with pdfplumber.open(filepath) as pdf:
                return "\n".join(page.extract_text() or '' for page in pdf.pages)
        elif filepath.endswith(".docx"):
            return "\n".join([para.text for para in Document(filepath).paragraphs])
        elif filepath.endswith(".doc"):
            print(f"[跳过] 不支持 .doc 文件: {filepath}")
            return ''
        elif filepath.endswith((".html", ".htm", ".shtml")):
            with open(filepath, "rb") as f:
                raw_data = f.read()
                encoding = chardet.detect(raw_data)["encoding"] or "utf-8"
                text = raw_data.decode(encoding, errors="ignore")
                soup = BeautifulSoup(text, "html.parser")
                return soup.get_text(separator='\n', strip=True)
        else:
            print(f"[跳过] 不支持的文件类型: {filepath}")
            return ''
    except Exception as e:
        print(f"[读取失败] {filepath} 错误: {str(e)}")
        return ''


import random
import os
def select_files_by_category(folder_path, target_entity):
    """
    从每个子文件夹中随机选择一个文件，并返回文件路径和类别信息

    参数:
    - folder_path: 主文件夹路径（如：D:\data\Auto_recommendtion\周鹏1）
    - target_entity: 目标实体名（如：周鹏）

    返回:
    - list: 包含(文件路径, 类别名)的元组列表
    """
    selected_files = []

    if not os.path.exists(folder_path):
        print(f"❌ 文件夹不存在: {folder_path}")
        return selected_files

    # 获取所有子文件夹
    subdirs = [d for d in os.listdir(folder_path)
               if os.path.isdir(os.path.join(folder_path, d))]

    print(f"📁 找到 {len(subdirs)} 个子文件夹: {subdirs}")

    for subdir in subdirs:
        subdir_path = os.path.join(folder_path, subdir)

        # 获取该文件夹中的所有文件
        all_files = [f for f in os.listdir(subdir_path)
                     if os.path.isfile(os.path.join(subdir_path, f))]

        if not all_files:
            print(f"⚠️  文件夹 '{subdir}' 中没有文件，跳过")
            continue

        # 随机选择一个文件
        selected_file = random.choice(all_files)
        selected_file_path = os.path.join(subdir_path, selected_file)

        selected_files.append((selected_file_path, subdir))
        print(f"✅ 选择文件: {subdir}/{selected_file}")

    return selected_files


def read_file_with_strategy(file_path, max_length=2500):
    """
    根据文件类型和长度智能读取文件内容

    参数:
    - file_path: 文件路径
    - max_length: 最大读取长度

    返回:
    - str: 文件内容
    """
    try:
        file_size = os.path.getsize(file_path)
        file_ext = os.path.splitext(file_path)[1].lower()

        print(f"📄 读取文件: {os.path.basename(file_path)} (大小: {file_size} bytes, 类型: {file_ext})")

        # 文本文件处理
        if file_ext in ['.txt', '.html', '.htm', '.xml', '.json']:
            return read_text_file(file_path, max_length)

        # PDF文件处理
        elif file_ext == '.pdf':
            return read_pdf_file(file_path, max_length)

        # Word文档处理
        elif file_ext in ['.docx', '.doc']:
            return read_word_file(file_path, max_length)

        # 其他文件类型
        else:
            print(f"⚠️  不支持的文件类型: {file_ext}，尝试作为文本文件读取")
            return read_text_file(file_path, max_length)

    except Exception as e:
        print(f"❌ 读取文件 {os.path.basename(file_path)} 失败: {e}")
        return ""


def read_text_file(file_path, max_length):
    """读取文本文件"""
    try:
        # 尝试多种编码
        encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    content = f.read()

                    # 如果文件太大，截取前一半或最大长度
                    if len(content) > max_length:
                        # 取前一半或最大长度，取较小值
                        half_length = min(len(content) // 2, max_length)
                        content = content[:half_length]
                        print(f"📏 文本过长，截取前 {half_length} 字符")

                    return content
            except UnicodeDecodeError:
                continue
        return ""
    except Exception as e:
        print(f"❌ 读取文本文件失败: {e}")
        return ""


def read_pdf_file(file_path, max_length):
    """读取PDF文件"""
    try:
        with pdfplumber.open(file_path) as pdf:
            total_text = ""
            page_count = min(len(pdf.pages), 5)  # 最多读取5页

            for i in range(page_count):
                page_text = pdf.pages[i].extract_text() or ''
                total_text += page_text + "\n"

                # 如果已经达到最大长度，停止读取
                if len(total_text) >= max_length:
                    total_text = total_text[:max_length]
                    print(f"📏 PDF文件过大，截取前 {max_length} 字符")
                    break

            return total_text
    except Exception as e:
        print(f"❌ 读取PDF文件失败: {e}")
        return ""


def read_word_file(file_path, max_length):
    """读取Word文档"""
    try:
        doc = Document(file_path)
        total_text = ""

        for para in doc.paragraphs:
            if para.text.strip():
                total_text += para.text + "\n"

                # 如果已经达到最大长度，停止读取
                if len(total_text) >= max_length:
                    total_text = total_text[:max_length]
                    print(f"📏 Word文档过大，截取前 {max_length} 字符")
                    break

        return total_text
    except Exception as e:
        print(f"❌ 读取Word文档失败: {e}")
        return read_text_file(file_path, max_length)  # 失败时尝试作为文本文件读取


def generate_prompt(person_name: str, raw_text: str, max_keywords: int = 15) -> str:
    """
    原有的专业提示词生成函数（保持兼容性）
    """
    return f"""你是一个信息检索专家，请从以下原始文本中，**逐字提取出**最多 {max_keywords} 个与人物 [{person_name}] 高度相关、便于网络搜索的关键词短语。

# 请严格遵守以下约束：

1. 所有关键词必须**逐字出现在原文中**，不能编造、概括或扩展；
2. 优先提取以下两类内容：
   - **机构/单位**：如"清华大学""自动化系"
   - **较大的研究方向**：如"人工智能""生物医学""计算机科学"
3. 如果原文中某些机构为长结构（如"安徽大学计算机科学与技术学院"），请**合理拆分为多个关键词**，例如：
   - "安徽大学"
   - "计算机科学与技术学院"

# 输出格式：
- **只输出关键词，不要输出任何解释、说明或其他文字**
- 中文分号分隔（使用中文分号"；"）
- 每条关键词应具有独立的搜索价值
- 所有关键词必须来自原文，且不超过 {max_keywords} 个
- 如果原文中确实没有相关关键词，请输出：**NO_KEYWORDS_FOUND**

原始文本如下：
{raw_text}
"""


def validate_and_extract_keywords(model_output, person_name):
    """
    验证并提取模型输出的关键词

    返回:
    - valid_keywords: 有效的关键词列表
    - is_valid: 是否成功提取到有效关键词
    """
    if not model_output or not model_output.strip():
        return [], False

    # 检查是否没有找到关键词的特殊标记
    if "NO_KEYWORDS_FOUND" in model_output:
        return [], True  # 明确表示没有关键词

    # 检查是否包含解释性文字（模型没有遵守指令）
    explanation_indicators = [
        "根据提供的原始文本", "未找到", "未提及", "建议检查", "文本内容主要",
        "输出为空", "抱歉", "无法提取", "建议提供", "确认"
    ]

    for indicator in explanation_indicators:
        if indicator in model_output:
            print(f"⚠️  检测到模型输出解释性文字: {indicator}")
            return [], False

    # 尝试按分号分割关键词
    keywords = [kw.strip() for kw in re.split(r"[；;]", model_output) if kw.strip()]

    # 验证关键词质量
    valid_keywords = []
    for keyword in keywords:
        # 过滤掉过短或无意义的关键词
        if (len(keyword) >= 2 and  # 至少2个字符
                not keyword.isdigit() and  # 不是纯数字
                not keyword.startswith(('http://', 'https://')) and  # 不是URL
                keyword != person_name):  # 不是实体名本身
            valid_keywords.append(keyword)

    # 检查是否提取到有效关键词
    if not valid_keywords:
        print("⚠️  未提取到有效关键词")
        return [], False

    print(f"✅ 验证通过的关键词: {valid_keywords}")
    return valid_keywords, True


def call_big_model(text: str, person_name: str, max_keywords: int, api_key: str) -> Optional[str]:
    """
    调用大模型 API，支持新的智能提示格式
    """
    API_URL = "https://api.siliconflow.cn/v1/chat/completions"
    MODEL_NAME = "deepseek-ai/DeepSeek-V3"
    HEADERS = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json"
    }

    # 检查是否是智能提示格式（包含类别信息）
    if "=== " in text and "类别文档内容 ===" in text:
        # 已经是智能提示格式，直接使用
        prompt = text
    else:
        # 使用原有的提示词格式
        prompt = generate_prompt(person_name, text, max_keywords)

    payload = {
        "model": MODEL_NAME,
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.1,
        "max_tokens": 200,
        "stop": ["\n\n"]
    }

    max_retries = 2
    for attempt in range(max_retries):
        try:
            print(f"🔄 API调用尝试 {attempt + 1}/{max_retries}")
            response = requests.post(API_URL, headers=HEADERS, json=payload, timeout=30)
            response.raise_for_status()

            result = response.json()['choices'][0]['message']['content'].strip()
            print(f"✅ API调用成功，返回长度: {len(result)}")

            return result

        except Exception as e:
            print(f"❌ API请求失败 ({attempt + 1}): {e}")
            if attempt < max_retries - 1:
                time.sleep(1)

    return None


def merge_all_texts(file_paths: list, max_workers: int = 10) -> str:
    """
    直接合并文件列表内容（跳过文件夹扫描）
    """

    def read_file(path):
        try:
            # 检测文件编码
            with open(path, 'rb') as f:
                raw_data = f.read()
                encoding = chardet.detect(raw_data)['encoding'] or 'utf-8'

            # 用检测到的编码读取文件
            with open(path, 'r', encoding=encoding, errors='ignore') as f:
                content = f.read().strip()
                if content:
                    return path, content
                else:
                    print(f"⚠️ 文件内容为空: {os.path.basename(path)}")
                    return path, None

        except UnicodeDecodeError:
            print(f"⚠️ 编码问题: {os.path.basename(path)}，尝试其他编码")
            # 尝试常见编码
            for enc in ['gbk', 'gb2312', 'latin-1', 'iso-8859-1']:
                try:
                    with open(path, 'r', encoding=enc, errors='ignore') as f:
                        content = f.read().strip()
                        if content:
                            print(f"✅ 使用 {enc} 编码成功读取: {os.path.basename(path)}")
                            return path, content
                except:
                    continue
            print(f"❌ 无法读取文件: {os.path.basename(path)}")
            return path, None
        except Exception as e:
            print(f"❌ 读取失败 {os.path.basename(path)}: {str(e)}")
            return path, None

    valid_contents = []
    with ThreadPoolExecutor(max_workers=max_workers) as executor:
        futures = {executor.submit(read_file, path): path for path in file_paths}
        for future in as_completed(futures):
            path, content = future.result()
            if content:
                valid_contents.append(content)

    return "\n\n".join(valid_contents).strip()


def clean_entity_name(entity: str) -> str:
    """
    自动去除实体名末尾的数字
    示例:
        "吴信东1" -> "吴信东"
        "实体123" -> "实体"
        "无数字" -> "无数字"
    """
    return re.sub(r'\d+$', '', entity)


def build_intelligent_prompt(selected_files, target_entity, max_keywords=15):
    """
    构建智能提示文本，整合原有的专业提示词格式

    参数:
    - selected_files: 选择的文件列表，包含(文件路径, 类别名)
    - target_entity: 目标实体名
    - max_keywords: 最大关键词数量

    返回:
    - str: 构建好的提示文本
    """
    # 构建内容部分
    content_parts = []
    total_length = 0
    max_total_length = 20000  # 总长度限制

    for file_path, category in selected_files:
        if total_length >= max_total_length:
            print("⚠️  内容部分过长，停止添加更多文件")
            break

        # 读取文件内容
        content = read_file_with_strategy(file_path, max_length=2500)  # 每个文件最多2500字符

        if not content.strip():
            continue

        # 添加类别信息
        category_header = f"=== {category}类别文档内容 ===\n"
        file_content = category_header + content + "\n\n"

        # 检查是否超过总长度限制
        if total_length + len(file_content) > max_total_length:
            remaining = max_total_length - total_length
            if remaining > 100:  # 至少保留100字符
                file_content = file_content[:remaining]
                print(f"📏 截断内容，保留前 {remaining} 字符")
            else:
                break

        content_parts.append(file_content)
        total_length += len(file_content)

    # 如果没有提取到任何内容
    if not content_parts:
        return generate_prompt(target_entity, "无相关文档内容", max_keywords)

    # 合并所有内容
    all_content = "".join(content_parts)

    # 使用原有的专业提示词格式
    prompt = f"""你是一个信息检索专家，请从以下不同类别的文档内容中，**逐字提取出**最多 {max_keywords} 个与人物 [{target_entity}] 高度相关、便于网络搜索的关键词短语。

# 请严格遵守以下约束：

1. 所有关键词必须**逐字出现在原文中**，不能编造、概括或扩展；
2. 优先提取以下两类内容：
   - **机构/单位**：如"清华大学""自动化系"
   - **较大的研究方向**：如"人工智能""生物医学""计算机科学"
3. 如果原文中某些机构为长结构（如"安徽大学计算机科学与技术学院"），请**合理拆分为多个关键词**，例如：
   - "安徽大学"
   - "计算机科学与技术学院"
4. **仔细分析每个类别的文档**，确保不遗漏任何重要信息

# 输出格式：
- **只输出关键词，不要输出任何解释、说明或其他文字**
- 中文分号分隔（使用中文分号"；"）
- 每条关键词应具有独立的搜索价值
- 所有关键词必须来自原文，且不超过 {max_keywords} 个
- 如果所有文档中确实没有相关关键词，请输出：**NO_KEYWORDS_FOUND**

# 文档来源说明：
以下内容来自该人物的多个相关文档，按类别分组显示：
{all_content}

"""

    print(f"📋 构建的提示文本总长度: {len(prompt)} 字符")
    return prompt


def filter_search_results_by_keywords(search_result, keywords):
    """
    根据关键词过滤搜索结果
    :param search_result: 搜索结果字典，包含data字段
    :param keywords: 关键词列表，至少需要2个关键词
    :return: 过滤后的搜索结果
    """
    if not keywords or len(keywords) < 2:
        print("⚠️ 关键词不足，跳过过滤")
        return search_result

    if not isinstance(search_result, dict) or 'data' not in search_result:
        print("⚠️ 搜索结果格式不正确，跳过过滤")
        return search_result

    filtered_data = []
    first_keyword = keywords[0]
    second_keyword = keywords[1]

    print(f"🔍 开始关键词过滤: 第一关键词='{first_keyword}', 第二关键词='{second_keyword}'")

    for item in search_result['data']:
        # 检查title或content是否包含第一个关键词
        title_contains_first = first_keyword in item.get('title', '')
        content_contains_first = first_keyword in item.get('content', '')

        # 如果title或content包含第一个关键词，再检查是否包含第二个关键词
        if title_contains_first or content_contains_first:
            title_contains_second = second_keyword in item.get('title', '')
            content_contains_second = second_keyword in item.get('content', '')

            if title_contains_second or content_contains_second:
                filtered_data.append(item)

    print(f"✅ 关键词过滤完成: 从 {len(search_result['data'])} 个结果中筛选出 {len(filtered_data)} 个相关结果")

    # 更新搜索结果
    search_result['data'] = filtered_data
    search_result['filtered_count'] = len(filtered_data)
    search_result['filter_keywords'] = [first_keyword, second_keyword]

    return search_result



# def auto_recommendtion(request: HttpRequest) -> dict:
#     """
#     主推荐函数 - 使用智能文件选择策略
#     """
#     name = request.GET.get("name", "").strip()
#     user_id = request.GET.get("user_id", "").strip()
#     api_key = request.GET.get("api_key", "sk-uwokmhxknecolbmvcrnfstfrcqzjeuekvnxfoghzrakeqybw")
#     max_keywords = int(request.GET.get("max_keywords", 5))
#     num_pages_to_crawl = int(request.GET.get("num_pages_to_crawl", 120))
#
#     print(f"🔍 开始处理: name={name}, user_id={user_id}")
#
#     if not name or not user_id:
#         return {"status": "error", "message": "Missing required parameters: name or user_id"}
#
#     try:
#         # 1. 下载文件
#         print("📥 步骤1: 下载相关文件...")
#         folder = test(name, int(user_id))
#         base_dir = os.path.join('D:\\data\\Auto_recommendtion', name)
#         print(f"✅ 下载完成，文件保存在: {base_dir}")
#
#         # 2. 智能选择文件
#         print("📁 步骤2: 从每个类别中随机选择文件...")
#         selected_files = select_files_by_category(base_dir, clean_entity_name(name))
#
#         if not selected_files:
#             print("❌ 没有找到可用的文件，使用实体名作为默认查询")
#             combined_query = clean_entity_name(name)
#             keywords = []
#             keywords_status = "no_files"
#         else:
#             print(f"✅ 从 {len(selected_files)} 个类别中选择的文件")
#
#             # 3. 构建智能提示文本
#             print("📝 步骤3: 构建智能提示文本...")
#             intelligent_text = build_intelligent_prompt(selected_files, clean_entity_name(name), max_keywords)
#
#             if not intelligent_text.strip():
#                 print("⚠️ 未能提取到有效文本内容")
#                 combined_query = clean_entity_name(name)
#                 keywords = []
#                 keywords_status = "no_content"
#             else:
#                 # 4. 提取关键词
#                 print("🤖 步骤4: 调用大模型提取关键词...")
#                 print(f"📋 发送给模型的文本长度: {len(intelligent_text)} 字符")
#
#                 start_time = time.time()
#                 result = call_big_model(
#                     text=intelligent_text,
#                     person_name=clean_entity_name(name),
#                     max_keywords=max_keywords,
#                     api_key=api_key
#                 )
#                 end_time = time.time()
#                 print(f"⏱️ 关键词提取耗时: {end_time - start_time:.2f}秒")
#
#                 if result:
#                     print(f"🔑 模型返回结果: {result}")
#
#                     # 验证和提取关键词
#                     keywords, is_valid = validate_and_extract_keywords(result, clean_entity_name(name))
#
#                     if is_valid and keywords:
#                         print(f"✅ 提取到的有效关键词: {keywords}")
#                         if len(keywords) >= 2:
#                             combined_query = f"{clean_entity_name(name)} {keywords[0]}"#{keywords[1]}
#                             keywords_status = "valid_multiple"
#                         elif len(keywords) == 1:
#                             combined_query = f"{clean_entity_name(name)} {keywords[0]}"
#                             keywords_status = "valid_single"
#                         else:
#                             combined_query = clean_entity_name(name)
#                             keywords_status = "no_keywords_found"
#                     else:
#                         print("❌ 模型返回无效结果，使用实体名")
#                         combined_query = clean_entity_name(name)
#                         keywords = []
#                         keywords_status = "invalid_output"
#                 else:
#                     print("❌ 模型调用失败，使用实体名")
#                     combined_query = clean_entity_name(name)
#                     keywords = []
#                     keywords_status = "api_failure"
#
#         print(f"🔍 最终搜索查询: {combined_query} (状态: {keywords_status})")
#
#         # 5. 进行网络搜索
#         print("🌐 步骤5: 进行网络搜索...")
#         mock_request = type('MockRequest', (), {'GET': {
#             "name": combined_query,
#             "num_pages_to_crawl": num_pages_to_crawl,
#             "userID": user_id,
#             "xiaoqi_name": clean_entity_name(name),
#             "enable_deduplication": "true"
#         }})
#
#         search_result = search_urls(mock_request)
#         print(
#             f"✅ 搜索完成，找到 {len(search_result.get('data', [])) if isinstance(search_result, dict) else len(search_result)} 个结果")
#
#         print("🔍 步骤6: 进行关键词过滤...")
#         combined_query_list = [clean_entity_name(name)] + keywords
#         search_result = filter_search_results_by_keywords(search_result, combined_query_list)
#
#
#
#         return {
#             "status": "success",
#             "data": search_result,
#             "keywords": keywords,
#             "combined_query": combined_query,
#             "file_count": len(selected_files),
#             "keywords_status": keywords_status,
#             "selected_categories": [category for _, category in selected_files] if selected_files else []
#         }
#
#     except Exception as e:
#         print(f"❌ 发生异常: {str(e)}")
#         import traceback
#         traceback.print_exc()
#         return {"status": "error", "message": str(e)}
#

def get_keywords_from_db(db, entity_name, user_id):
    """
    从数据库获取实体关键词
    """
    # 直接查询 key_words 字段
    query = """
    SELECT key_words 
    FROM xiaoqi_new 
    WHERE xiaoqi_name = %s
    LIMIT 1
    """

    try:
        with db.connection.cursor() as cursor:
            cursor.execute(query, (entity_name,))
            result = cursor.fetchone()

            if not result:
                print("❌ 数据库中未找到对应实体记录")
                return [], False

            key_words_value = result[0] if result else None

            if not key_words_value:
                print("⚠️ 关键词字段为空")
                return [], False

            print(f"🔍 原始关键词数据: {key_words_value} (类型: {type(key_words_value)})")

            # 解析关键词
            keywords = []

            if isinstance(key_words_value, str):
                if key_words_value.startswith('[') and key_words_value.endswith(']'):
                    # JSON 数组格式
                    try:
                        keywords = json.loads(key_words_value)
                        print(f"✅ 解析为JSON数组: {keywords}")
                    except json.JSONDecodeError as e:
                        print(f"❌ JSON解析失败: {e}")
                        # 尝试分号分隔
                        keywords = [kw.strip() for kw in key_words_value.split(';') if kw.strip()]
                else:
                    # 分号分隔或普通字符串
                    keywords = [kw.strip() for kw in key_words_value.split(';') if kw.strip()]
                    print(f"✅ 解析为分号分隔: {keywords}")

            elif isinstance(key_words_value, (list, tuple)):
                # 已经是列表格式
                keywords = list(key_words_value)
                print(f"✅ 直接使用列表: {keywords}")

            if keywords:
                print(f"✅ 从数据库解析到关键词: {keywords}")
                return keywords, True
            else:
                print("⚠️ 关键词字段有值但解析后为空")
                return [], False

    except pymysql.MySQLError as e:
        print(f"查询关键词失败：{e}")
        return [], False
    except Exception as e:
        print(f"解析关键词失败：{e}")
        return [], False

def check_table_structure(db, table_name="xiaoqi_new"):
    """
    检查表结构
    """
    try:
        with db.connection.cursor() as cursor:
            cursor.execute(f"DESCRIBE {table_name}")
            result = cursor.fetchall()
            print(f"🔍 [DEBUG] 表结构 {table_name}:")
            for row in result:
                print(f"  {row}")
    except Exception as e:
        print(f"❌ 检查表结构失败: {e}")



def update_keywords_to_db(db, entity_name, user_id, keywords):
    """
    更新关键词到数据库 - 添加详细调试
    """
    print(f"🔍 [DEBUG] update_keywords_to_db 开始执行")
    print(f"🔍 [DEBUG] 参数: entity_name={entity_name}, user_id={user_id}")
    print(f"🔍 [DEBUG] keywords={keywords}")
    print(f"🔍 [DEBUG] keywords类型: {type(keywords)}")

    check_query = """
    SELECT COUNT(*) 
    FROM xiaoqi_new 
    WHERE xiaoqi_name = %s
    """

    update_query = """
    UPDATE xiaoqi_new 
    SET key_words = %s 
    WHERE xiaoqi_name = %s
    """

    insert_query = """
    INSERT INTO xiaoqi_new (xiaoqi_name, key_words) 
    VALUES (%s, %s)
    """

    try:
        with db.connection.cursor() as cursor:
            # 检查记录是否存在
            print(f"🔍 [DEBUG] 执行检查查询: {check_query} with {entity_name}")
            cursor.execute(check_query, (entity_name,))
            exists = cursor.fetchone()[0] > 0
            print(f"🔍 [DEBUG] 记录是否存在: {exists}")

            # 准备关键词数据
            if isinstance(keywords, list):
                keywords_json = json.dumps(keywords, ensure_ascii=False)
                print(f"🔍 [DEBUG] 列表转换为JSON: {keywords_json}")
            else:
                keywords_json = str(keywords)
                print(f"🔍 [DEBUG] 非列表直接转为字符串: {keywords_json}")

            print(f"🔍 [DEBUG] 最终要保存的数据: {keywords_json}")
            print(f"🔍 [DEBUG] 数据类型: {type(keywords_json)}")

            if exists:
                print(f"🔍 [DEBUG] 执行更新: {update_query} with ({keywords_json}, {entity_name})")
                cursor.execute(update_query, (keywords_json, entity_name))
                print(f"✅ 更新现有记录: {entity_name}")
            else:
                print(f"🔍 [DEBUG] 执行插入: {insert_query} with ({entity_name}, {keywords_json})")
                cursor.execute(insert_query, (entity_name, keywords_json))
                print(f"✅ 插入新记录: {entity_name}")

            print(f"🔍 [DEBUG] 提交事务")
            db.connection.commit()
            print(f"✅ 关键词已成功{'更新' if exists else '插入'}到数据库")
            return True

    except Exception as e:
        print(f"❌ 更新关键词失败：{e}")
        print(f"🔍 [DEBUG] 错误类型: {type(e)}")
        print(f"🔍 [DEBUG] 完整错误信息:")
        import traceback
        traceback.print_exc()
        db.connection.rollback()
        return False

def auto_recommendtion(request: HttpRequest) -> dict:
    """
    主推荐函数 - 优化版：优先从数据库读取关键词
    """
    name = request.GET.get("name", "").strip()
    user_id = request.GET.get("user_id", "").strip()
    api_key = request.GET.get("api_key", "sk-uwokmhxknecolbmvcrnfstfrcqzjeuekvnxfoghzrakeqybw")
    max_keywords = int(request.GET.get("max_keywords", 5))
    num_pages_to_crawl = int(request.GET.get("num_pages_to_crawl", 40))

    print(f"🔍 开始处理: name={name}, user_id={user_id}")


    if not name or not user_id:
        return {"status": "error", "message": "Missing required parameters: name or user_id"}

    # 初始化数据库连接
    db = None
    try:
        db = MySQLDatabase(
            host="114.213.234.179",
            user="koroot",
            password="DMiC-4092",
            database="db_hp"
        )
        db.connect()

        print("🔍 [DEBUG] 检查数据库表结构...")
        check_table_structure(db)

        # 1. 首先尝试从数据库获取关键词
        clean_name = clean_entity_name(name)
        print("📋 步骤1: 从数据库查询关键词...")
        db_keywords, has_keywords = get_keywords_from_db(db, name, user_id)

        if has_keywords and db_keywords:
            print(f"✅ 从数据库获取到关键词: {db_keywords}")
            keywords = db_keywords
            keywords_source = "database"
            keywords_status = "from_db"

            # 构建搜索查询
            if len(keywords) >= 2:
                combined_query = f"{clean_name} {keywords[0]}"  # 使用第一个关键词
            elif len(keywords) == 1:
                combined_query = f"{clean_name} {keywords[0]}"
            else:
                combined_query = clean_name

        else:
            print("❌ 数据库中没有关键词，开始文件处理流程...")
            keywords_source = "file_analysis"

            # 2. 下载文件
            print("📥 步骤2: 下载相关文件...")
            folder = test(name, int(user_id))
            base_dir = os.path.join('D:\\data\\Auto_recommendtion', name)
            print(f"✅ 下载完成，文件保存在: {base_dir}")

            # 3. 智能选择文件
            print("📁 步骤3: 从每个类别中随机选择文件...")
            selected_files = select_files_by_category(base_dir, clean_name)

            if not selected_files:
                print("❌ 没有找到可用的文件，使用实体名作为默认查询")
                combined_query = clean_name
                keywords = []
                keywords_status = "no_files"
            else:
                print(f"✅ 从 {len(selected_files)} 个类别中选择的文件")

                # 4. 构建智能提示文本
                print("📝 步骤4: 构建智能提示文本...")
                intelligent_text = build_intelligent_prompt(selected_files, clean_name, max_keywords)

                if not intelligent_text.strip():
                    print("⚠️ 未能提取到有效文本内容")
                    combined_query = clean_name
                    keywords = []
                    keywords_status = "no_content"
                else:
                    # 5. 提取关键词
                    print("🤖 步骤5: 调用大模型提取关键词...")
                    print(f"📋 发送给模型的文本长度: {len(intelligent_text)} 字符")

                    start_time = time.time()
                    result = call_big_model(
                        text=intelligent_text,
                        person_name=clean_name,
                        max_keywords=max_keywords,
                        api_key=api_key
                    )
                    end_time = time.time()
                    print(f"⏱️ 关键词提取耗时: {end_time - start_time:.2f}秒")

                    if result:
                        print(f"🔑 模型返回结果: {result}")

                        # 验证和提取关键词
                        keywords, is_valid = validate_and_extract_keywords(result, clean_name)

                        if is_valid and keywords:
                            print(f"✅ 提取到的有效关键词: {keywords}")

                            # 在调用 update_keywords_to_db 之前添加
                            print(f"🔍 [DEBUG] 准备调用 update_keywords_to_db")
                            print(f"🔍 [DEBUG] 当前关键词: {keywords}")
                            print(f"🔍 [DEBUG] 关键词类型: {type(keywords)}")
                            print(f"🔍 [DEBUG] 是否是列表: {isinstance(keywords, list)}")
                            if isinstance(keywords, list):
                                print(f"🔍 [DEBUG] 列表长度: {len(keywords)}")
                                print(f"🔍 [DEBUG] 列表内容: {keywords}")

                            # 6. 将新关键词保存到数据库
                            print("💾 步骤6: 将关键词保存到数据库...")
                            update_success = update_keywords_to_db(db, name, int(user_id), keywords)
                            if update_success:
                                print("✅ 关键词已成功保存到数据库")
                            else:
                                print("⚠️ 关键词保存到数据库失败")

                            if len(keywords) >= 2:
                                combined_query = f"{clean_name} {keywords[0]}"
                                keywords_status = "valid_multiple"
                            elif len(keywords) == 1:
                                combined_query = f"{clean_name} {keywords[0]}"
                                keywords_status = "valid_single"
                            else:
                                combined_query = clean_name
                                keywords_status = "no_keywords_found"
                        else:
                            print("❌ 模型返回无效结果，使用实体名")
                            combined_query = clean_name
                            keywords = []
                            keywords_status = "invalid_output"
                    else:
                        print("❌ 模型调用失败，使用实体名")
                        combined_query = clean_name
                        keywords = []
                        keywords_status = "api_failure"

        print(f"🔍 最终搜索查询: {combined_query} (来源: {keywords_source}, 状态: {keywords_status})")

        # 7. 进行网络搜索
        print("🌐 步骤7: 进行网络搜索...")
        mock_request = type('MockRequest', (), {'GET': {
            "name": combined_query,
            "num_pages_to_crawl": num_pages_to_crawl,
            "userID": user_id,
            "xiaoqi_name": clean_name,
            "enable_deduplication": "true"
        }})

        search_result = search_urls(mock_request)
        result_count = len(search_result.get('data', [])) if isinstance(search_result,
                                                                        dict) and 'data' in search_result else 0
        print(f"✅ 搜索完成，找到 {result_count} 个结果")

        # 8. 进行关键词过滤
        print("🔍 步骤8: 进行关键词过滤...")
        combined_query_list = [clean_name] + keywords
        search_result = filter_search_results_by_keywords(search_result, combined_query_list)

        return {
            "status": "success",
            "data": search_result,
            "keywords": keywords,
            "combined_query": combined_query,
            "keywords_source": keywords_source,
            "keywords_status": keywords_status,
            "from_database": has_keywords if keywords_source == "database" else False
        }

    except Exception as e:
        print(f"❌ 发生异常: {str(e)}")
        import traceback
        traceback.print_exc()
        return {"status": "error", "message": str(e)}
    finally:
        # 关闭数据库连接
        if db and db.connection:
            db.close()

#
# def debug_full_process():
#     """逐步调试完整流程"""
#     print("=== 完整流程调试 ===")
#
#     # 模拟请求
#     class MockRequest:
#         def __init__(self):
#             self.GET = {
#                 "name": "屠呦呦1",
#                 "user_id": "6000622",
#                 "max_keywords": "5",
#                 "num_pages_to_crawl": "20"
#             }
#
#     request = MockRequest()
#     result = auto_recommendtion(request)
#
#     print(f"\n=== 最终结果 ===")
#     print(f"状态: {result.get('status')}")
#     print(f"关键词: {result.get('keywords', [])}")
#     print(f"搜索查询: {result.get('combined_query', '')}")
#     print(f"文件数量: {result.get('file_count', 0)}")
#     print(f"完整返回: {result}")
#
#
#
# if __name__ == "__main__":
#     # print("=== 测试文件读取 ===")
#     # test_file_reading()
#     #
#     # print("\n=== 测试模型API ===")
#     # test_model_api()
#     #
#     # print("\n=== 调试完整流程 ===")
#     T0= time.time()
#     debug_full_process()
#     T1 = time.time()
#     print("全部运行时间",T1-T0)
#
